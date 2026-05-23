"""
Append Session 35 (May 6, 2026) to the running session report.

Mirrors the structure used in earlier sessions: a date/title line, a Branch line,
an Executive Summary, then numbered sections with prose paragraphs (no bullet
gimmicks), a Files Summary, Git Commits, and an end-of-report marker.
"""

from pathlib import Path
from docx import Document

DOC = Path(r"F:/infoWebWorld/InfoWebWorld_Session_Report_2026-03-23.docx")

doc = Document(str(DOC))


def H(text: str) -> None:
    """Plain paragraph used as a section heading line in the existing doc style."""
    doc.add_paragraph(text)


def P(text: str) -> None:
    """Plain paragraph (body)."""
    doc.add_paragraph(text)


# ───────────────────────────────────────────────────────────────────────────
# Header
# ───────────────────────────────────────────────────────────────────────────
H("Session 35 — May 6, 2026")
H(
    "Session: Live Listing Page Parity with /test-listing-page — 11 Sections "
    "Reworked, Sidebar Alternatives Bug Fix, Integrations Rich-Object Migration"
)
H("Developer: Claude Code (AI) + Aadil Parmar")
H("Branch: main (production)")

H("Executive Summary")
P(
    "A focused session driven entirely by Aadil's frustration that the live "
    "/company/[slug] page was diverging from the polished /test-listing-page "
    "design — sections vanishing on real listings, layout bleeding, and "
    "Mailchimp sample data leaking into real renders. Eleven sections were "
    "audited, rewritten, and unified so that the live page now matches the "
    "test-page UI for every datum the owner can populate from /dashboard/new — "
    "with graceful empty-state cards whenever a field is missing instead of "
    "the section disappearing or rendering bare placeholder text. The "
    "biggest single change was migrating the integrations field from a flat "
    "string[] to a rich {name, website?, description?}[] schema, with the "
    "form Step 5 rebuilt around repeating rows and the live page rendering "
    "favicon-from-website logos plus owner-supplied descriptions. Two real "
    "bugs were fixed along the way: the sidebar mini-Alternatives card had "
    "no CSS for its <img> logo so a screenshot URL would render at native "
    "1200x800 and blow the card open with overflow text, and the All "
    "Features matrix was falling back to Mailchimp's 53-row sample whenever "
    "a real listing had key_features populated but no flat features[]. "
    "Build clean, TypeScript clean, single commit pushed to main "
    "(484bf86) and live via Vercel auto-deploy."
)

# ───────────────────────────────────────────────────────────────────────────
# 1. Audit: How real mode was diverging from preview
# ───────────────────────────────────────────────────────────────────────────
H("1. Audit — How Real Mode Was Diverging from /test-listing-page")
P(
    "Before any code changes, the entire ListingDetailPage.tsx (3,164 lines) "
    "was grepped for every isPreview gate. Findings fell into three buckets. "
    "Bucket A — sections gated entirely by hasReviews or isPreview that "
    "vanished on a real listing with no reviews yet (Review Insights, the "
    "User Opinions block inside Pricing). Bucket B — sections that DID "
    "render in real mode but used a totally different, simpler layout from "
    "preview (Customer Support's right-only options card, Alternatives' "
    "small sib-cards instead of rich alt-cards, Who Uses' bare ul lists "
    "instead of donut+diamond+bars, FAQs collapsing to nothing). Bucket C — "
    "the worst kind: real listings unintentionally rendering preview sample "
    "data because of non-null fallback chains. The most egregious case was "
    "filteredFeatures defaulting to ALL_FEATURES (Mailchimp's 53-row sample) "
    "whenever a real listing didn't have a features[] array — even though "
    "the section was visible because key_features WAS set. /company/"
    "gemini-ai-715bb01d was rendering real Gemini AI key features at the top "
    "and Mailchimp's "
    "Email Templates / Drag and Drop / Audience Segmentation list directly "
    "below it. Aadil saw this and rightly flagged it as broken."
)

# ───────────────────────────────────────────────────────────────────────────
# 2. Sidebar Alternatives mini — overflow bug
# ───────────────────────────────────────────────────────────────────────────
H("2. Sidebar Alternatives Mini Card — Overflow Bug Fix")
P(
    "The bug: a screenshot of the sidebar showed a sibling listing's image "
    "(a laptop with a QR-code generator) rendering at full natural size with "
    "the company name 'test' overflowing on the right edge, breaking out of "
    "the card. Root cause: in test-listing-page.css, the .tlp-alt-mini-logo "
    "letter-fallback class had width/height/border-radius rules but the "
    ".tlp-alt-mini-logo-img class (used when an actual <img> was rendered) "
    "had ZERO CSS rules at all. When the sibling's logo_url pointed to a "
    "screenshot URL or an oversized PNG, the <img> rendered at the source "
    "image's natural dimensions (could be 1200x800), pushing the card open "
    "and overflowing the rest of the sidebar."
)
P(
    "Fix shipped: added .tlp-alt-mini-logo-img with width/height 40x40, "
    "border-radius 10px, object-fit: contain, light bg #F9FAFB, 1px subtle "
    "border, 4px padding, flex-shrink: 0. Also bumped the letter fallback "
    "from 36x36 circle to 40x40 rounded-square so both variants share visual "
    "rhythm. Truncation rules added to .tlp-alt-mini-name (white-space "
    "nowrap, overflow hidden, text-overflow ellipsis). The hijack of "
    ".tlp-alt-mini-rate (with inline-style overrides) for the meta line was "
    "replaced with a proper .tlp-alt-mini-meta class. Real mode meta line "
    "now reads category name plus 'From $X period' (whichever fields exist) "
    "with truncated tagline as a final fallback. Logo favicon size bumped "
    "from 64 to 128 for crispness at the 40px display size. Preview Brevo "
    "card with Stars + rating was left untouched."
)

# ───────────────────────────────────────────────────────────────────────────
# 3. The "fake data" rule — why the All Features bug was the worst
# ───────────────────────────────────────────────────────────────────────────
H("3. The Fake-Data Rule and the All Features Bug")
P(
    "Aadil's stated principle drove every decision: 'the info giving is in "
    "the hand of user but to display them in the best way is in ours so if "
    "there is less info still it should not happen bro.' Translation: the "
    "owner controls what they share, but the platform controls how that data "
    "renders. Missing data should mean a graceful empty state, never a "
    "section that vanishes (the page looks broken) and never a fallback to "
    "another listing's data (Mailchimp leaks)."
)
P(
    "The All Features fake-data bug: filteredFeatures (a useMemo around line "
    "1463-1472) returned realFeatures.map(...) when present, otherwise "
    "ALL_FEATURES (the 53-row Mailchimp sample). That sample existed for the "
    "design-preview workflow on /test-listing-page. But the section was "
    "gated by (realKeyFeatures || realFeatures || isPreview). So a real "
    "listing with key_features populated and no flat features[] would render "
    "the section AND fall through to the Mailchimp sample. Fix: the useMemo "
    "now returns [] when not isPreview and no realFeatures. The section's "
    "render branch checks filteredFeatures.length and switches to a "
    ".tlp-empty-card with a magnifier-list icon and copy 'Full feature list "
    "coming soon — {Co} hasn't shared the complete breakdown yet.' The "
    "search input is hidden when there are fewer than 9 features. Same "
    "pattern applied to Key Features: source priority is realKeyFeatures "
    "(rich, name+description) → first 6 of realFeatures mapped to "
    "name-only → empty card with a star icon. The preview-only 'Based on "
    "our analysis of 984 reviews' intro paragraph stays preview-only; in "
    "real mode a soft intro replaces it: 'Top features {Co} highlights "
    "about themselves — what the team considers most important about the "
    "product.'"
)

# ───────────────────────────────────────────────────────────────────────────
# 4. Alternatives — rich card grid in real mode
# ───────────────────────────────────────────────────────────────────────────
H("4. Alternatives — Rich Card Grid in Real Mode")
P(
    "Before: real-mode Alternatives rendered as small .tlp-sib-card boxes "
    "with logo + name + tagline + footer (price + 'View →'). Preview mode "
    "rendered the rich .tlp-alt-grid with logo, rating row, 'Learn More' "
    "CTA, 'Starting from' price block with currency split, free trial / "
    "version checkmark flags, and a four-attribute detailed ratings list "
    "(Ease of Use, Features, Value for Money, Customer Support). The two "
    "looked almost nothing alike."
)
P(
    "After: real mode renders the same .tlp-alt-grid layout but populated "
    "ONLY with real sibling fields. Per card: logo (s.logo_url, falls back "
    "to clearbit favicon from s.website domain at size 128, falls back to a "
    "letter tile via new .tlp-alt-logo--letter modifier), company name, "
    "category sub-line, 'Learn More' CTA pointing to /company/[slug], "
    "'Starting from' price block (parses s.starting_price as numeric, "
    "renders with $ symbol; muted clock icon + 'Pricing not shared' when "
    "absent), and a 3-line clamped tagline (.tlp-alt-tagline with "
    "-webkit-line-clamp 3). Dropped: the rating row (no aggregated review "
    "data per sibling), the free trial / version flags (we don't capture "
    "those as booleans on the sibling rows server-side), and the detailed "
    "ratings list (no NLP-derived per-attribute scores). New CSS classes "
    "added at end of test-listing-page.css: .tlp-alt-logo--letter (teal bg, "
    "white letter), .tlp-alt-sub (mute, 12px, ellipsis), .tlp-alt-tagline "
    "(3-line clamp)."
)

# ───────────────────────────────────────────────────────────────────────────
# 5. Who Uses — donut/diamond/bars from real data
# ───────────────────────────────────────────────────────────────────────────
H("5. Who Uses — Render Donut, Diamond Cluster, Bars from Real Data")
P(
    "Before: real mode rendered three bare ul lists labeled 'Target company "
    "sizes', 'Industries served', 'Use cases' — pure text with bullets. "
    "Preview rendered the rich 3-column grid: vertical capsule bars for "
    "company size (with hardcoded percentages), teal-shade SVG donut for "
    "industries (hover reveals percentage tooltip), and a 5-shape diamond "
    "cluster for use cases. Aadil specifically called out the graphs not "
    "rendering and asked that they appear when the owner provides data."
)
P(
    "After: a single IIFE inside the !isPreview block renders the same "
    ".tlp-wu-grid 3-column layout from real data. Industries become an "
    "EQUAL-SHARE donut — each industry is allocated 1/N of the circumference "
    "with one color from a teal palette of eight values "
    "(#006B6B → #DFF5F5). Hover shows the industry NAME only — no fake "
    "percentages, no fake review counts. Use cases render as the existing "
    "diamond shape positions (max 5, color-coded from the same teal palette) "
    "plus a full legend list below for any overflow. Company size renders "
    "the existing 3-bar layout: Small Businesses / Midsize Businesses / "
    "Enterprises. Bars matching a string in realCompanySizes (case-"
    "insensitive prefix match against 'Small', 'Mid', 'Enterprise' to handle "
    "'Small (1-50)' etc.) get full 100% height; unpicked bars get height "
    "12% and class .is-faded (opacity .35) — visually unmistakable that the "
    "owner picked the highlighted ones without inventing sentiment data. "
    "Per-column fallback: when an array is empty, a small "
    ".tlp-wu-empty 'Not specified yet' dashed card sits in that column. "
    "The section gate was kept (any-data || isPreview) — if the owner has "
    "shared none of the three, the whole section stays hidden rather than "
    "rendering three empty placeholders."
)

# ───────────────────────────────────────────────────────────────────────────
# 6. Review Insights — always render with empty-state CTA
# ───────────────────────────────────────────────────────────────────────────
H("6. Review Insights — Always Render with Empty-State Write-Review CTA")
P(
    "Before: section was gated (hasReviews || isPreview). On a real listing "
    "with zero reviews, the whole #insights section vanished. Real mode also "
    "used a bare 'User reviews' header + a flat list of reviews — totally "
    "unlike preview's rich 2-column grid (avatar stack + sentiment bar + "
    "topic chips on the left, analyzed paragraph + quote cards on the right)."
)
P(
    "After: gate dropped — section always renders. Real mode now uses the "
    "same .tlp-in-grid 2-column layout. Left aside has Overall rating block "
    "(real overallRating + Stars + count when hasReviews; '—' + 'No reviews "
    "yet' when none) plus a Topic insights placeholder block ('We'll surface "
    "trending topics and sentiment once {Co} collects more reviews.'). Right "
    "main column shows real review cards as quotes when reviewsData.recent "
    "exists — using a new .tlp-in-quote-title class for the review title, "
    ".tlp-in-quote-av--img modifier for the avatar img variant, .tlp-in-"
    "quote-date for the dot-separated date — OR a polished .tlp-in-empty "
    "card (chat-bubble SVG, headline 'No reviews yet for {Co}', sub-copy, "
    "and a coral pill 'Write the first review' button that triggers the "
    "existing review modal via the same setReviewOpen / requireLogin auth "
    "gate). NLP-derived sentiment bars and topic chips were dropped from "
    "real mode entirely — they need analysis we don't run."
)

# ───────────────────────────────────────────────────────────────────────────
# 7. User Opinions on Price and Value — render in real mode
# ───────────────────────────────────────────────────────────────────────────
H("7. User Opinions on Price and Value — Render in Real Mode")
P(
    "This block lived inside the Pricing section as a preview-only "
    "(isPreview &&) closure showing 'Value for money rating: 4.4 (17.5K)' + "
    "a 4-card grid of fake review quotes (Mailchimp sample). Aadil "
    "explicitly listed it as missing from live listings."
)
P(
    "Added a parallel !isPreview branch immediately after. Same H3 'User "
    "opinions about {Co} price and value', same .tlp-price-meta value-for-"
    "money rating row (real overallRating + count when hasReviews, muted "
    "'No ratings yet' otherwise — using a new .tlp-vo-norating class). When "
    "reviews exist: real reviewsData.recent.slice(0, 4) rendered through the "
    "existing .tlp-vo card style, with avatar img variant via "
    ".tlp-vo-av--img and inline Stars rating below the name. When no "
    "reviews exist: a .tlp-vo-empty card (star SVG, headline 'What do you "
    "think of {Co} pricing?', sub-copy, coral 'Write a review' pill button) "
    "— same auth gate as the Insights empty CTA."
)

# ───────────────────────────────────────────────────────────────────────────
# 8. Customer Support — full 2-col layout in real mode
# ───────────────────────────────────────────────────────────────────────────
H("8. Customer Support — Full 2-Column Layout in Real Mode")
P(
    "Before: section was gated (realSupportChannels || realTrainingOptions || "
    "isPreview). When neither was provided the whole section vanished. Real "
    "mode rendered ONLY the right-side .tlp-cs-opts card; the left analysis "
    "panel (heading + rating + sentiment bullets + intro) was preview-only. "
    "Bottom 3-quote .tlp-cs-quotes block was preview-only too."
)
P(
    "After: section gate dropped — always renders. Real mode left column now "
    "shows: H3 'What do users say about {Co} customer support?' + customer "
    "support rating row (real Stars + count when hasReviews, 'No ratings "
    "yet' muted span otherwise) + a soft contextual line that switches "
    "between 'See review snippets below to learn what users say...' (when "
    "reviews exist) and '{Co} hasn't collected enough reviews to surface "
    "support insights yet — write the first one to share your experience' "
    "(otherwise). The SUPPORT_BULLETS narrative bullets were dropped from "
    "real mode (NLP-derived). Right column still renders Support options + "
    "Training options cards from realSupportChannels / realTrainingOptions, "
    "and when neither is provided shows a new .tlp-cs-opts-empty dashed "
    "card 'Support & training options coming soon — {Co} hasn't shared "
    "their channels yet.' Below the grid, real mode renders up to 3 review "
    "snippets as .tlp-cs-quote cards (with a new .tlp-cs-quote-av--img "
    "modifier for avatar images and Stars below name) when hasReviews; "
    "otherwise a .tlp-vo-empty card 'How was your experience with {Co} "
    "support?' with a 'Write a review' coral pill (re-using the value-"
    "opinions empty-state styles)."
)

# ───────────────────────────────────────────────────────────────────────────
# 9. FAQs — always render with empty-state card
# ───────────────────────────────────────────────────────────────────────────
H("9. FAQs — Always Render with Empty-State Card")
P(
    "Before: gated by (realFaqs || isPreview). Real listings with no FAQs "
    "had the whole section disappear."
)
P(
    "After: gate dropped. Section header + lead always render. The IIFE "
    "that builds the items array now returns realFaqs (when present), "
    "FAQS sample (when isPreview), or [] otherwise. When items is empty, "
    "a .tlp-empty-card renders with a question-mark-circle SVG and copy "
    "'FAQs coming soon — {Co} hasn't added frequently asked questions "
    "yet. Check back soon — or reach out with your question.' The 'reach "
    "out' phrase is a /contact link styled with .tlp-inline-link. The "
    "accordion render path is unchanged when items exist."
)

# ───────────────────────────────────────────────────────────────────────────
# 10. Integrations — rich-object schema migration
# ───────────────────────────────────────────────────────────────────────────
H("10. Integrations — Rich-Object Schema Migration")
P(
    "Aadil's specific complaint: 'these in dashboard i am talking from where "
    "the users list the listing they have to give proper website url and "
    "write something per integration as it is in the mailchip.' The flat "
    "string[] integrations field was insufficient — the test page renders "
    "rich integration cards with logo + name + description quote, and we "
    "couldn't replicate that without per-integration metadata."
)
P(
    "Schema. New IntegrationItem type {name: string; website?: string; "
    "description?: string} added in two places: app/iww-hq/data/submissions-"
    "storage.ts (server-side type) and app/dashboard/new/form/types.ts "
    "(form state type, with non-optional website/description because the "
    "form always controls those input strings). RealSubmission.integrations "
    "changed from string[] to IntegrationItem[]. FormState.integrations "
    "changed similarly. The DB column (submissions.integrations) remains a "
    "JSON column — the shape change is forward-compatible and required no "
    "ALTER TABLE."
)
P(
    "Tolerance for legacy data. The mapRow function in submissions-storage.ts "
    "and the local mapServerRow in ListingDetailPage.tsx both run the "
    "incoming JSON array through a normalizer: each item is wrapped as "
    "{name: string} if it's a plain string, mapped to "
    "{name, website?, description?} if it's an object, dropped otherwise. "
    "Empty-name entries are filtered out. This means existing production "
    "rows with integrations: ['Stripe', 'Slack'] continue rendering "
    "correctly as plain chips while new submissions get the rich UI."
)
P(
    "Form rebuild. Step 5 (Step5Features.tsx) had its integrations field "
    "rewritten. The single ChipInput was replaced with a repeating-row UI "
    "matching the existing key-features pattern: each row is a "
    ".df-int-row CSS grid (3-col on desktop: name / website URL / remove + "
    "description spanning the bottom; reflows to 4-row stack at "
    "max-width 600px) with a name input (maxLength 60), a url input "
    "(type=url, maxLength 200), a 2-row description textarea "
    "(maxLength 400), and a remove icon button. Add button '+ Add "
    "integration' caps at MAX_INTEGRATIONS = 24. Helper functions "
    "updateIntegration, addIntegration, removeIntegration mirror the "
    "key-feature helpers. CSS for .df-int-row added in dashboard-form.css "
    "right after .df-kf-row."
)
P(
    "Submit handler. DashboardListingForm.tsx's submit payload now maps the "
    "form rows to clean objects via filter(name.trim()).map(...) — emits "
    "{name} alone when website/description are empty, omits the optional "
    "fields rather than storing empty strings. The /api/submissions POST "
    "handler unchanged: it JSON.stringifies whatever body.integrations "
    "array it receives, which already supports objects."
)
P(
    "Live page render. ListingDetailPage's integrations section now splits "
    "view.realIntegrations into 'rich' (has website OR description) and "
    "'simple' (just name). Rich entries render as .tlp-int-card cards in "
    "the existing 2-col .tlp-int-grid: logo from clearbit(domain) when "
    "website is provided (or a teal letter tile via "
    ".tlp-int-logo--letter), name in the existing .tlp-int-name style, an "
    "external-link line with the bare domain via new .tlp-int-link, and "
    "the description in a new .tlp-int-desc paragraph. Plain-name entries "
    "fall through to a small 'Also integrates with' header (.tlp-int-also, "
    "uppercase mute label) and the existing .tlp-int-chips chip list. The "
    "preview-only 'Integrations rated by users' chrome (rating cards, "
    "author quotes, pager arrows) stays preview-only — those are NLP-"
    "derived and we don't have that data."
)
P(
    "Other call sites. The overview Q&A 'most popular integrations' line "
    "was updated to .map(i => i.name).join(', '). Step 7 review summary "
    "row similarly maps to names. The /sector search keyword pool in "
    "SectorHero.tsx was updated to .map(i => i.name) so integration names "
    "remain searchable. The admin /iww-hq/submissions detail card "
    "renders {t.name} per chip instead of the raw string."
)

# ───────────────────────────────────────────────────────────────────────────
# 11. Files modified
# ───────────────────────────────────────────────────────────────────────────
H("11. Files Modified (10 files)")
P(
    "app/listing/ListingDetailPage.tsx — every section rewrite plus the "
    "mapServerRow integrations normalizer and the overview Q&A name-mapping "
    "fix. By far the largest change: 841 lines net delta."
)
P(
    "app/styles/test-listing-page.css — appended ~250 lines of new classes "
    "across the empty-state vocabulary (.tlp-empty-card / -ico / -body / "
    "-title / -sub), Who Uses (.tlp-wu-bar.is-faded, .tlp-wu-empty), "
    "Insights (.tlp-in-rate-num--muted, .tlp-in-coming, .tlp-in-quote-title, "
    ".tlp-in-quote-av--img, .tlp-in-quote-date, .tlp-in-empty plus -ico/"
    "-title/-sub/-cta), Value Opinions (.tlp-vo-norating, .tlp-vo-av--img, "
    ".tlp-vo-empty plus -ico/-body/-title/-sub/-cta), Customer Support "
    "(.tlp-cs-opts-empty, .tlp-cs-quote-av--img, .tlp-cs-rate em), "
    "Integrations (.tlp-int-logo--letter, .tlp-int-link, .tlp-int-desc, "
    ".tlp-int-also), Alternatives (.tlp-alt-logo--letter, .tlp-alt-sub, "
    ".tlp-alt-tagline), and the sidebar mini fix (.tlp-alt-mini-logo-img, "
    ".tlp-alt-mini-meta plus name truncation)."
)
P(
    "app/iww-hq/data/submissions-storage.ts — IntegrationItem type added, "
    "RealSubmission.integrations changed to IntegrationItem[], mapRow "
    "rewritten with the legacy-string-tolerant normalizer."
)
P(
    "app/dashboard/new/form/types.ts — IntegrationItem type added, "
    "FormState.integrations changed from string[] to IntegrationItem[]."
)
P(
    "app/dashboard/new/form/steps/Step5Features.tsx — ChipInput replaced "
    "with the repeating row UI plus helper functions and MAX_INTEGRATIONS "
    "constant."
)
P(
    "app/dashboard/new/form/DashboardListingForm.tsx — submit payload "
    "integrations mapping rewritten."
)
P(
    "app/dashboard/new/form/steps/Step7Review.tsx — integrations row maps "
    "to .name."
)
P(
    "app/styles/dashboard-form.css — .df-int-row grid + responsive override "
    "appended after .df-kf-row."
)
P(
    "app/sector/components/SectorHero.tsx — search keyword pool maps "
    "integrations to .name."
)
P(
    "app/iww-hq/submissions/page.tsx — admin chip render uses {t.name}."
)

# ───────────────────────────────────────────────────────────────────────────
# 12. Build + Git
# ───────────────────────────────────────────────────────────────────────────
H("12. Build & Git")
P(
    "TypeScript clean (./node_modules/.bin/tsc --noEmit, exit 0). "
    "npm run build clean (Next.js 16 Turbopack, ✓ Compiled successfully in "
    "~4s, all 90 static pages generated). Committed as 484bf86 'feat: live "
    "listing parity with /test-listing-page across 11 sections' with a "
    "detailed body documenting every section change. Pushed to "
    "origin/main; Vercel auto-deploys main → infowebworld.com. Diff "
    "stats: 10 files changed, 1,001 insertions, 211 deletions."
)

# ───────────────────────────────────────────────────────────────────────────
# 13. Decisions captured
# ───────────────────────────────────────────────────────────────────────────
H("13. Decisions Captured")
P(
    "Equal-share donut over fake percentages. The real Industries donut "
    "could have rendered with hardcoded sample percentages or with even "
    "1/N slices. Chose equal-share — each industry the owner picked gets "
    "an equal slice — because that visualizes a real fact ('these are the "
    "industries served') without claiming review-derived market share."
)
P(
    "Faded bars over hidden bars for company size. Could have rendered "
    "only the picked sizes. Chose to render all three with unpicked at "
    "12% height + 35% opacity so the visual structure stays intact and "
    "the contrast between picked and unpicked is the data signal."
)
P(
    "Empty-state cards over hidden sections. Major sections (Insights, "
    "Customer Support, FAQs, Key Features, All Features) now ALWAYS "
    "render with empty cards when data is missing. The Pros & Cons "
    "sidebar block stays gated because it's a sidebar accessory and the "
    "main column flows correctly without it. The Who Uses section stays "
    "gated when ALL three sub-fields are empty, because rendering three "
    "empty placeholders side-by-side reads worse than the section being "
    "absent."
)
P(
    "No DB migration for integrations rich schema. The integrations column "
    "is JSON; legacy string[] payloads are tolerated by the normalizer; "
    "new submissions write rich objects. Both shapes coexist in the same "
    "column without a backfill or ALTER TABLE."
)
P(
    "NLP-derived chrome stripped from real mode everywhere. Sentiment "
    "bars, topic chips, narrative bullets ('users report X'), per-feature "
    "ratings, integration ratings, integration author quotes — all "
    "preview-only because we don't run review NLP and won't fake it. When "
    "we add real review aggregation in the future, those blocks can be "
    "re-enabled with real data; until then, they don't render in real "
    "mode."
)

H("— End of Session 35 Report —")

doc.save(str(DOC))
print(f"Saved Session 35 to {DOC}")

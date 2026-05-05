# -*- coding: utf-8 -*-
"""
Append a detailed Session 34 entry to the InfoWebWorld session-report .docx.

Mirrors the pattern from write-session30.py — adds a page break, then
heading hierarchy + paragraphs + bullets via python-docx.
"""
from docx import Document
from docx.shared import Pt, RGBColor

doc = Document(r'F:\infoWebWorld\InfoWebWorld_Session_Report_2026-03-23.docx')


def add_h(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def add_p(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    if bold:
        run.bold = True


def add_b(text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'


doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
add_h('Session 34 \u2014 May 5, 2026', 1)
add_p('Session: Listing Form V3 \u2014 Full Rebuild to Capture Every Datum the Live '
      'Listing Page Renders, 18 New DB Columns, Listing-Page Fallback Wiring, Demo-Data '
      'Helper, Sticky-Footer Bleed Fix', bold=True)
add_p('Developer: Claude Code (AI) + Aadil Parmar')
add_p('Branch: main (production) \u2014 same commit as Session 33 deliverables')

add_h('Executive Summary', 2)
add_p('A long, surgical session focused entirely on the /dashboard/new listing form. '
      'Aadil pushed back on the existing form\u2019s UI (\u201cI don\u2019t like the UI of '
      'the current listing form at all\u201d) and gave a clear directive: redesign it to '
      'match the compact pill geometry of the dashboard sidebar (no glow, no gradients, '
      'simple color), and capture every datum the live ListingDetailPage renders so newly-'
      'approved listings show real data instead of falling back to Mailchimp sample content. '
      'He also explicitly granted DB schema flexibility (\u201cI can edit the db so be '
      'flexible\u201d) so missing columns could be added rather than working around them.')
add_p('The session delivered: (1) a top-to-bottom audit of the 2,597-line ListingDetailPage '
      'inventorying every field, (2) a database migration adding 18 new optional columns, '
      '(3) full propagation through the API, RealSubmission type, mapServerRow + view object '
      'inside the listing page, and per-section render fallbacks so old listings keep the '
      'sample look while new submissions show real data, (4) a complete from-scratch form at '
      'app/dashboard/new/form/ with a 7-step pill-rail UI in the dashboard\u2019s exact '
      'visual vocabulary, (5) deletion of the old form code (ListingFormV2 + the entire '
      'business/form/ tree + listing-v2.css), (6) a \u201cFill demo data\u201d button that '
      'seeds every form field with realistic sample values for fast testing, and (7) a '
      'structural fix to the sticky-footer bleed where form content was visible past the '
      'sticky bottom bar. TypeScript and production build both verified clean. All work was '
      'pushed to main alongside the Session 33 deliverables that had been sitting '
      'uncommitted on the working tree.')

# ═══════════════════════════════════════════════════════════════════════════════
add_h('1. The Audit \u2014 Mapping Every Listing-Page Datum to a DB Field', 2)

add_p('The first hour was spent reading every section of app/listing/ListingDetailPage.tsx '
      '(2,597 lines) and cataloguing every visible datum. The page renders a sticky identity '
      'header, breadcrumb, page-title block, overview Q&A card, UI screenshots carousel, '
      'reviews-insights section with topic chips and quote cards, who-uses panel with '
      'company-size bars + industries donut + use-cases diamond cluster, key-features list, '
      'all-features matrix, alternatives grid, pricing plans, value-quotes, integrations, '
      'customer-support panel, FAQs accordion, popular comparisons, customers-also-viewed, '
      'and related categories.')

add_h('1.1 What Was Already DB-Backed (Before This Session)', 3)
add_b('Identity panel: companyName, logoUrl, tagline, hqLocation, phone, email, website, '
      'founded, employees \u2014 wired since Session 29.')
add_b('Overview Q&A: description, integrations (just for the integrations Q&A line).')
add_b('Stats column: founded, employees, realPricing.length.')
add_b('Pricing tiers: pricingTiers JSON column drives the pricing card grid.')
add_b('FAQs: faqs JSON column.')
add_b('Screenshots field exists in DB but the carousel rendered hardcoded Unsplash images.')
add_b('Features field exists in DB but the all-features matrix rendered a hardcoded sample.')

add_h('1.2 What Was Hardcoded Sample Data (Needed New DB Columns)', 3)
add_p('Every section below was rendering a hardcoded constant defined at the top of '
      'ListingDetailPage.tsx:')
add_b('HEADER_TAGS \u2014 5 short labels under the company name in the sticky header '
      '(Email Marketing, SaaS, Newsletters, etc.).')
add_b('PROS_SHORT (3 labels) and CONS_SHORT (3 labels) \u2014 the overview side card.')
add_b('Sidebar starting price \u2014 hardcoded \u201c$13/month\u201d with two checkmarks for '
      'free trial / free version.')
add_b('Support channels and training options \u2014 hardcoded six-channel list (Email/help '
      'desk, Chat, Knowledge base, FAQs/forum, 24/7 live rep, Phone support) + four-option '
      'training list.')
add_b('KEY_FEATURES \u2014 6 features with rich paragraph descriptions. Currently the '
      'features field on submissions is just string[], so even when populated it can\u2019t '
      'drive this section.')
add_b('Industries served (donut chart input), use cases (diamond cluster), target company '
      'sizes (3-bar chart) \u2014 the entire \u201cWho uses\u201d section.')
add_b('Languages supported, mobile app availability flags, compliance / certifications, '
      'awards / press \u2014 referenced in the listing copy or implied by the design.')

add_h('1.3 What Stays Hardcoded Even After This Session (Community Data)', 3)
add_p('Some data the submitter cannot supply, since it would be community-generated or '
      'platform-derived. These remain hardcoded sample values, will eventually be replaced '
      'by real systems (reviews, analytics) rather than by form input:')
add_b('Overall star rating + review count.')
add_b('Topic chips with sentiment (NLP-derived from reviews).')
add_b('Reviewer quotes (insight quotes, value quotes, support quotes).')
add_b('Per-feature ratings + review counts.')
add_b('Industry donut numeric percentages (analytics-derived).')
add_b('Per-integration rich quote cards.')
add_b('Alternative companies, popular comparisons, customers-also-viewed (auto-derived '
      'from category siblings in a future session).')
add_b('1\u20132 / 3\u20134 / 5-star sentiment-bar breakdown.')

# ═══════════════════════════════════════════════════════════════════════════════
add_h('2. Database Migration \u2014 18 New Optional Columns', 2)

add_p('File created: database/migration-listings-v3.sql. Idempotent-ish (use ADD COLUMN; '
      'if a column already exists MySQL errors with \u201cDuplicate column name\u201d \u2014 '
      'just skip that line). All columns optional (DEFAULT NULL or 0 for booleans). Safe to '
      'run while the site is live; existing INSERTs are not affected because new columns '
      'have defaults.')

add_h('2.1 The 18 New Columns', 3)
add_b('header_tags JSON \u2014 sticky-header category pills (3\u20135 short labels).')
add_b('pros JSON \u2014 3 short positive labels for the overview side card.')
add_b('cons JSON \u2014 3 short negative labels.')
add_b('industries_served JSON \u2014 donut-chart input (array of industry names).')
add_b('use_cases JSON \u2014 diamond-cluster input (array of use-case names).')
add_b('target_company_sizes JSON \u2014 array from {small, mid, enterprise}.')
add_b('key_features JSON \u2014 array of {name, description} for the rich top-features '
      'section. Adds editorial paragraphs to a feature, distinct from the flat features[].')
add_b('starting_price DECIMAL(10,2) \u2014 the \u201cFrom $X\u201d sidebar value.')
add_b('starting_price_period VARCHAR(20) \u2014 \u201c/ month\u201d, \u201c/ year\u201d, '
      '\u201cone-time\u201d, etc.')
add_b('has_free_trial TINYINT(1) \u2014 toggles the \u201cFree trial\u201d sidebar checkmark.')
add_b('has_free_version TINYINT(1) \u2014 toggles the \u201cFree version\u201d checkmark.')
add_b('support_channels JSON \u2014 multi-select from {email, chat, knowledge base, FAQs, '
      '24/7, phone}.')
add_b('training_options JSON \u2014 multi-select from {live, videos, webinars, docs, '
      'in-person}.')
add_b('languages JSON \u2014 array of language names (free-add allowed).')
add_b('has_ios_app TINYINT(1) and has_android_app TINYINT(1) \u2014 mobile app availability '
      'flags.')
add_b('compliance JSON \u2014 array of standards (GDPR, CAN-SPAM, SOC 2, HIPAA, ISO 27001, '
      'PCI-DSS, CCPA, FERPA + custom).')
add_b('awards JSON \u2014 array of {name, year} (e.g., \u201cG2 Leader 2026\u201d).')

add_h('2.2 Why JSON for Multi-Value Fields', 3)
add_p('Each multi-value field (header_tags, pros, cons, industries, use_cases, etc.) could '
      'have been modeled as a separate joined table (submission_industries, '
      'submission_compliance, etc.). JSON columns were chosen because: (1) the data is '
      'always read as a complete bag with the parent row \u2014 there\u2019s no foreign-key '
      'integrity benefit, (2) zero new tables means zero new joins on every listing-page '
      'read, (3) MySQL 8 native JSON support handles the storage efficiently, (4) the '
      'existing pattern (features, integrations, screenshots, faqs, pricing_tiers) is '
      'already JSON \u2014 staying consistent.')

# ═══════════════════════════════════════════════════════════════════════════════
add_h('3. API + Type Layer Wiring', 2)

add_h('3.1 RealSubmission Type Extended', 3)
add_p('app/iww-hq/data/submissions-storage.ts:')
add_b('Added type aliases KeyFeature = {name, description} and Award = {name, year?}.')
add_b('Added 18 new fields to RealSubmission: headerTags, pros, cons, industriesServed, '
      'useCases, targetCompanySizes, keyFeatures, startingPrice, startingPricePeriod, '
      'hasFreeTrial, hasFreeVersion, supportChannels, trainingOptions, languages, '
      'hasIosApp, hasAndroidApp, compliance, awards.')
add_b('Extended PricingTier type with optional features?: string[] so each tier can list '
      'its included features (the listing page renders these inside each pricing card).')
add_b('mapRow now reads all 18 new columns from the DB row, parsing JSON safely with the '
      'existing parseJson helper, converting tinyint(0/1) values to booleans for the flag '
      'columns, and stringifying starting_price for consistent display.')

add_h('3.2 POST /api/submissions Inserts the New Fields', 3)
add_p('app/api/submissions/route.ts:')
add_b('Added an arrJson() helper that JSON.stringifies a value only if it\u2019s a non-empty '
      'array, returning null otherwise. Keeps the DB column null for absent values rather '
      'than storing \u201c[]\u201d.')
add_b('Pulled all 18 new fields off the request body, normalised booleans to 0/1, normalised '
      'startingPrice to a Number (or null), and added them to the INSERT column list and '
      'parameters array.')
add_b('GET /api/submissions and /api/listings/[slug] both already use SELECT s.*, so the '
      'new columns flow through to clients automatically once mapRow knows about them \u2014 '
      'no edits needed in those endpoints.')

add_h('3.3 ListingDetailPage \u2014 view Object Extension + Render Fallbacks', 3)
add_p('app/listing/ListingDetailPage.tsx already had a `view` object that derives every '
      'identity field from real DB data with sample fallback. This session added 18 new '
      'view fields (realHeaderTags, realPros, realCons, realIndustries, realUseCases, '
      'realCompanySizes, realKeyFeatures, realStartingPrice, realStartingPeriod, '
      'realHasFreeTrial, realHasFreeVersion, realSupportChannels, realTrainingOptions, '
      'realLanguages, realHasIosApp, realHasAndroidApp, realCompliance, realAwards) and '
      'wired them into the render code with sample fallback. Specifically:')
add_b('Sticky-head category tags now render (view.realHeaderTags || HEADER_TAGS).')
add_b('Overview side-card pros/cons render (view.realPros || PROS_SHORT) and '
      '(view.realCons || CONS_SHORT).')
add_b('Sidebar starting price reads view.realStartingPrice / realStartingPeriod with '
      'fallback to the hardcoded $13/month.')
add_b('Free-trial / free-version checkmarks gated on the boolean flags.')
add_b('Customer-support \u201cSupport options\u201d list reads view.realSupportChannels '
      'with fallback to the hardcoded six-channel list.')
add_b('\u201cTraining options\u201d list reads view.realTrainingOptions with fallback to '
      'the hardcoded four-option list.')
add_b('UI screenshot carousel: const UI_IMAGES = view.realScreenshots || UI_IMAGES_FALLBACK '
      '\u2014 real screenshots from the DB now drive the carousel for the first time.')
add_b('Key-features section: when view.realKeyFeatures is present, renders submitter\u2019s '
      'rich features (name + description) with the rating star hidden (rating 0 \u2192 '
      'no star); when null, falls back to the hardcoded KEY_FEATURES sample (which has '
      'ratings + reviewer quotes).')
add_b('All-features matrix: when view.realFeatures is present, maps each feature name to '
      '[name, 0, 0] tuples and renders just the name (rating cluster hidden when count is '
      'zero); when null, falls back to the 53-row sample ALL_FEATURES.')

# ═══════════════════════════════════════════════════════════════════════════════
add_h('4. The New Form \u2014 app/dashboard/new/form/', 2)

add_h('4.1 Why a Full Rewrite, Not a Refactor', 3)
add_p('The existing form (app/business/ListingFormV2.tsx + app/business/form/) was built '
      'for the /business slide-in drawer in Session 27\u201328. After the drawer was '
      'removed in Session 29 and the form moved into the dashboard, the styling drifted '
      '\u2014 coral hero badges, gradient pricing chips, glowing focus rings, dashed-border '
      'image upload tiles with translate-on-hover \u2014 all of it conflicting with the '
      'flat, no-glow, dashboard-sidebar-pill aesthetic Aadil wanted. Refactoring the '
      'existing 1,331-line CSS file would be slower and riskier than starting fresh. '
      'Decision: build a parallel form scoped to the dashboard, verify, then delete the old.')

add_h('4.2 File Layout', 3)
add_p('New directory: app/dashboard/new/form/ (20 files).')
add_b('types.ts \u2014 FormState (45 fields), PlanCaps (with hasFaqs, hasKeyFeatures, '
      'hasPricingTiers, hasAudienceSection, hasComplianceAndAwards toggles), KeyFeature, '
      'Award, PricingTier (now with optional features[]), StepDef, StepProps + variants.')
add_b('constants.ts \u2014 INITIAL state, PLAN_CAPS for free / starter / yearly / lifetime '
      'with per-plan caps (maxScreenshots, maxFeatures, maxKeyFeatures, maxFaqs, maxTags, '
      'maxPricingTiers, maxAwards, maxLanguages, maxIndustries, maxUseCases) and feature '
      'flags. Plus shared option lists: COMPANY_SIZE, SUPPORT_CHANNELS, TRAINING_OPTIONS, '
      'COMPLIANCE, TEAM_SIZE, FUNDING_STAGE, PRICING_MODEL, PRICING_PERIOD, COMMON_LANGUAGES.')
add_b('validation.ts \u2014 per-step validators (identity, category, contact, story, '
      'features, pricing, review). Hard-validates required fields and plan caps; soft-allows '
      'optional fields.')
add_b('demo.ts \u2014 buildDemoForm(categories) helper returning a complete sample '
      'Partial<FormState>. Skips logoUrl + screenshots (user uploads). Auto-picks first '
      'L1/L2/L3 from the loaded categories.')
add_b('DashboardListingForm.tsx \u2014 thin orchestrator (~270 lines): state, draft '
      'autosave, validation, step navigation, submit. Mounts the rail + content panel.')
add_b('components/ \u2014 9 shared components (Field, Select, PillToggle, ChipInput, '
      'Uploader, PhoneRow, CategoryPicker, RailNav, Footer).')
add_b('steps/ \u2014 7 step components (Step1Identity, Step2Category, Step3Contact, '
      'Step4Story, Step5Features, Step6Pricing, Step7Review).')

add_h('4.3 The 7-Step Structure', 3)
add_p('Compressed from the old form\u2019s 9 steps. Logic: every step maps to a contiguous '
      'visual block on the listing page, so the user is always filling in something that '
      'will end up in a specific section.')
add_b('Step 1 \u2014 Identity: logo upload (single 96px tile), company name, tagline, '
      'website, header tags (3-5 chips), description (multi-paragraph textarea).')
add_b('Step 2 \u2014 Category: 3-level cascading picker (sector \u2192 category \u2192 '
      'subcategory) + specializations (min 2, plan-capped) + tag groups (one-tag-per-group '
      'requirement).')
add_b('Step 3 \u2014 Contact: contact name, email, phone (country picker + dial code + '
      'number), country / state / city dropdowns from country-state-city, HQ address '
      'free-text.')
add_b('Step 4 \u2014 Story & media: screenshots (multi-upload, plan-capped), demo video URL, '
      'founded year, team size, funding stage, industries served, use cases, target company '
      'sizes (multi-pill), languages (pill toggle with free-add), iOS/Android app '
      'checkboxes, compliance pills (free-add), awards (name + year repeating row), three '
      'social URLs.')
add_b('Step 5 \u2014 Features: full features list (chip input, plan-capped), key features '
      'rich blocks (name + description repeating, plan-gated), integrations (chip input), '
      'support channels (multi-pill), training options (multi-pill).')
add_b('Step 6 \u2014 Pricing & FAQ: starting price + period + free-trial/version checkboxes, '
      'pricing model dropdown, pricing tiers (name + price + period + features chip-input '
      'per tier, plan-gated), pros (3 chips) + cons (3 chips), FAQs (Q + A repeating, '
      'plan-gated).')
add_b('Step 7 \u2014 Review: identity preview card + per-step summary rows with '
      '\u201cEdit step\u201d jump links + final submit.')

add_h('4.4 Visual System \u2014 Matching the Dashboard Sidebar', 3)
add_p('Aadil\u2019s explicit constraint: \u201cdesign the best compact pill ui simple color '
      'no glow and all.\u201d Every visual decision was guided by the existing .tp-* '
      'sidebar geometry and the no-glow rule.')
add_b('Step rail (220px left): vertical pill list. Each pill 32px tall, 999px radius, 14px '
      'padding \u2014 identical geometry to .tp-nav-item. Inactive: transparent background, '
      '14px Inter weight 500 black. Active: cream-filled (#F5F2EF), weight 600, chevron-right '
      'at end. Done: numbered pill turns into a check glyph, text muted. Numbered (01\u201307) '
      'rather than icon-prefixed for cleaner look.')
add_b('Bottom of rail: a 4px progress bar + \u201cX% complete\u201d label.')
add_b('Inputs: 12px-radius rounded rectangles (NOT pills \u2014 pills reserved for buttons '
      'and nav). 1.5px solid border (var(--df-border) = #E5E7EB). Focus changes border-color '
      'to coral. Zero box-shadow on focus, ever.')
add_b('Labels: 11px uppercase Inter weight 700 with 0.08em letter-spacing. Required marker '
      'is a small coral asterisk after the label.')
add_b('Buttons: primary \u201cContinue\u201d / \u201cSubmit\u201d is a coral pill (999px) '
      'with white text, no shadow, no glow. Hover darkens. Secondary buttons are outline '
      'pills with 1.5px black border, transparent fill.')
add_b('Tag chips (pill-toggle, chip-input): 999px pills, 13px Inter. Selected = coral fill '
      '+ white text. Unselected = white + 1.5px border. No gradient, no glow.')
add_b('Image uploader: 1.5px solid (not dashed) border, 12px radius, plus-icon center. Drag '
      'state changes border to coral. Single-tile variant for logo, grid variant for '
      'screenshots.')
add_b('Plan-gated rows render visible-but-disabled with an inline \u201cUpgrade to '
      'unlock\u201d coral pill on the right of the label \u2014 better than hiding entire '
      'fields, builds upgrade signal.')
add_b('Inter font is enforced site-wide on the dashboard via dashboard-shell.css\u2019s '
      '!important rule on .tp-root descendants \u2014 inherited automatically by everything '
      'in the form.')

add_h('4.5 New CSS \u2014 app/styles/dashboard-form.css', 3)
add_p('700+ lines, .df-* namespace. Scoped under .tp-root so Inter inherits. Defines all of '
      ': layout grid, plan strip, rail nav (pill geometry), content card, section heads, '
      'fields, inputs, textareas, buttons, add-button, icon-button, layout helpers '
      '(grid-2 / grid-3 / empty / locked-preview), pill toggles, chip inputs, tag groups, '
      'category picker (numbered cascade), select (portal-based custom select), phone row '
      '(country + dial + number), uploader, checkbox row (custom-styled checkboxes that look '
      'like coral filled pills), award row, key-feature row, pricing-tier row, FAQ row, '
      'sticky footer, success state, review card, responsive breakpoints (1100, 900, 600).')

add_h('4.6 Submit Payload Wiring', 3)
add_p('DashboardListingForm.onSubmit re-validates every prior step, then POSTs to '
      '/api/submissions with the full FormState including all 18 new fields. The endpoint '
      'has been updated in this same session (section 3.2) to consume them. On success, '
      'localStorage draft cleared, success card shown with the future URL slug.')

# ═══════════════════════════════════════════════════════════════════════════════
add_h('5. \u201cFill Demo Data\u201d Button', 2)

add_p('After Aadil ran the form for the first time he asked for a button to seed the entire '
      'form with sample data so he could test the listing page without typing every field '
      'by hand. \u201cAdd a button demo in the listing clicked that add the demo data images '
      'and logo i will add manually.\u201d')

add_h('5.1 Implementation', 3)
add_b('New file: app/dashboard/new/form/demo.ts. Exports buildDemoForm(categories: '
      'Category[]): Partial<FormState>. Returns a complete sample state with realistic '
      'Mailchimp-style values for every field except logoUrl + screenshots (user uploads).')
add_b('Auto-picks the first L1 \u2192 first L2 of that L1 \u2192 first L3 of that L2 from '
      'the loaded categories. listingTypeIds and tagIds remain empty because they depend on '
      'l3Id and load asynchronously \u2014 user picks them manually after the rerender.')
add_b('Sample data covers all sections: identity (name, tagline, header tags, description), '
      'contact (name, email, phone, US/CA/SF), story (demo video, founded, team, funding, '
      'social URLs, app flags, compliance, awards, languages, industries, use cases, '
      'company sizes, support channels, training options), features (10 features, 3 rich '
      'key features, 10 integrations), pricing (starting price $13/month, 3 tiers with '
      'per-tier feature lists, 3 pros, 2 cons, 3 FAQs).')
add_b('Button rendered in the plan strip\u2019s right side (next to \u201cDraft saved\u201d): '
      'small flat 11px Inter 700 pill with a coral star icon + \u201cFill demo data\u201d '
      'label. Hover transitions to coral border + coral text + faint coral wash background.')
add_b('Click handler: setForm(prev => ({ ...prev, ...buildDemoForm(allCategories) })) \u2014 '
      'preserves any logo / screenshots already uploaded.')

# ═══════════════════════════════════════════════════════════════════════════════
add_h('6. Sticky Footer Bleed Fix', 2)

add_h('6.1 The Bug', 3)
add_p('Aadil reported a UI glitch: \u201con scroll the form goes below it basically form is '
      'visible from below also and it is not looking good ui glitch feels.\u201d The sticky '
      'footer (Step X of 7 + Continue) was visually being passed through by form content as '
      'the user scrolled the page.')

add_h('6.2 Root Cause', 3)
add_p('The footer was position: sticky; bottom: 0 inside .df-content (the form card). The '
      'actual scroll container was .tp-content (the dashboard main column \u2014 a flex '
      'child of .tp-main with overflow-y: auto). When content overflowed, .df-content '
      'extended below the viewport. The sticky footer pinned to the .tp-content viewport '
      'edge, NOT the .df-content card edge \u2014 so as the user scrolled, body content was '
      'visible in the gap between the sticky footer\u2019s position and the actual bottom '
      'border of the form card.')

add_h('6.3 The Fix \u2014 Make the Form Card Itself the Scroll Container', 3)
add_b('app/styles/dashboard-shell.css: added .tp-main:has(.df-wrap) > .tp-content { '
      'overflow: hidden; display: flex; flex-direction: column; min-height: 0 } \u2014 '
      'disables the dashboard\u2019s outer scroll only when the form is mounted '
      '(detected via :has() selector on the .df-wrap class).')
add_b('Same file: .tp-content > .nl--form { flex: 1; min-height: 0; display: flex; '
      'flex-direction: column } \u2014 lets the NewListingClient wrapper inherit the '
      'available height so the DashboardHeader takes its natural size and the form fills '
      'the rest.')
add_b('Mobile override (max-width 900px) restores the page-flow scroll model: .tp-content '
      'goes back to overflow: visible because the rail collapses to a horizontal strip on '
      'mobile and an internal-scroll card would feel cramped.')
add_b('app/styles/dashboard-form.css: .df-wrap, .df-grid \u2014 added flex: 1 + min-height: '
      '0 so they fill the available height inside .tp-content.')
add_b('.df-content \u2014 added overflow: hidden + display: flex + flex-direction: column. '
      'The card itself becomes the scroll boundary; everything inside is clipped to the '
      'rounded card edges, preventing any bleed.')
add_b('.df-body \u2014 changed to flex: 1 + min-height: 0 + overflow-y: auto + '
      'overscroll-behavior: contain. The body now scrolls inside the card, with a thin '
      '8px gray scrollbar styled to match.')
add_b('.df-footer \u2014 dropped position: sticky entirely. Replaced with flex-shrink: 0. '
      'The footer naturally pins to the bottom of .df-content because the body is the '
      'flex-grow child consuming the rest \u2014 no sticky needed, no bleed possible.')

add_h('6.4 Why This Pattern Is Right', 3)
add_p('The dashboard shell already uses the same :has() selector for the home overview '
      '(.tp-main:has(.dash-home) { overflow: hidden }) so the 7-tile dashboard fits a single '
      'viewport. Reusing the same pattern for the form keeps the codebase consistent and '
      'the responsive override identical. The footer-as-natural-flex-child pattern (rather '
      'than position: sticky) is also more robust because it does not depend on the scroll '
      'container being the immediate parent \u2014 it just sits at the bottom of the card '
      'because the body fills the rest.')

# ═══════════════════════════════════════════════════════════════════════════════
add_h('7. Cleanup \u2014 Old Form Code Deleted', 2)
add_p('After the new form was wired into NewListingClient and verified building, the old '
      'form code was deleted in one sweep:')
add_b('app/business/ListingFormV2.tsx (the orchestrator).')
add_b('app/business/ListingForm.tsx (the legacy V1 form, never used after V2 shipped but '
      'left around).')
add_b('app/business/form/ \u2014 entire directory: components/ (Field, CustomSelect, '
      'CategoryPicker, FlagImg, ImageUploader, PhoneRow, StepHead, TagPillSelector), '
      'steps/ (StepAbout, StepBusiness, StepCategory, StepContact, StepFaq, StepFeatures, '
      'StepLocation, StepPremium, StepReview), constants.ts, icons.tsx, types.ts, '
      'validation.ts.')
add_b('app/styles/listing-v2.css (1,334 lines of glow / gradient / hover-lift CSS).')
add_b('app/globals.css \u2014 dropped the listing-v2.css @import, added the new '
      'dashboard-form.css @import.')
add_b('app/business/GetListedLanding.tsx \u2014 PlanKey type import re-pointed to the new '
      'app/dashboard/new/form/DashboardListingForm.')

# ═══════════════════════════════════════════════════════════════════════════════
add_h('8. Build + TypeScript Verification', 2)
add_b('./node_modules/.bin/tsc --noEmit \u2014 clean (zero errors). One transient error '
      'caught and fixed mid-session: NewListingClient referenced c.hasPremium on PlanCaps '
      'which no longer exists \u2014 swapped to c.hasKeyFeatures + c.hasPricingTiers + '
      'c.hasComplianceAndAwards which are the new plan flags.')
add_b('npm run build \u2014 \u2713 Compiled successfully in ~4.5s, all 89 static pages '
      'generated without errors.')
add_b('Dev server (Next.js 16.2 Turbopack) hot-reloaded every edit cleanly. Final '
      'next-development.log entries: \u201c\u2713 Compiled in 60-77ms\u201d for the last '
      '5 hot reloads.')
add_b('Auth-gate verification: GET /dashboard/new \u2192 307 to /business for unauthed '
      'users, exactly as expected (the dashboard layout enforces requireDashboardUser()).')

# ═══════════════════════════════════════════════════════════════════════════════
add_h('9. Files Created (This Session)', 2)
add_b('database/migration-listings-v3.sql \u2014 18-column ALTER TABLE migration.')
add_b('app/dashboard/new/form/types.ts \u2014 FormState, PlanCaps + 4 helper types.')
add_b('app/dashboard/new/form/constants.ts \u2014 INITIAL, PLAN_CAPS, 9 option lists.')
add_b('app/dashboard/new/form/validation.ts \u2014 per-step validators.')
add_b('app/dashboard/new/form/demo.ts \u2014 buildDemoForm() helper.')
add_b('app/dashboard/new/form/DashboardListingForm.tsx \u2014 orchestrator (~270 lines).')
add_b('app/dashboard/new/form/components/Field.tsx \u2014 label + control wrapper with '
      'lockedReason support.')
add_b('app/dashboard/new/form/components/Select.tsx \u2014 portal-based custom select.')
add_b('app/dashboard/new/form/components/PillToggle.tsx \u2014 multi-select pill toggle '
      'with optional free-add input.')
add_b('app/dashboard/new/form/components/ChipInput.tsx \u2014 free-text chip list with '
      'Enter-to-add.')
add_b('app/dashboard/new/form/components/Uploader.tsx \u2014 image upload (single + grid '
      'variants).')
add_b('app/dashboard/new/form/components/PhoneRow.tsx \u2014 country picker + dial code + '
      'phone input.')
add_b('app/dashboard/new/form/components/CategoryPicker.tsx \u2014 numbered 3-level '
      'cascade.')
add_b('app/dashboard/new/form/components/RailNav.tsx \u2014 left rail pill nav + progress '
      'bar.')
add_b('app/dashboard/new/form/components/Footer.tsx \u2014 Back / counter / Continue/Submit '
      'row.')
add_b('app/dashboard/new/form/steps/Step1Identity.tsx through Step7Review.tsx \u2014 7 '
      'step components.')
add_b('app/styles/dashboard-form.css \u2014 700+ lines, .df-* namespace.')
add_b('scripts/write-session34.py \u2014 this script (the one writing the doc entry you '
      'are reading).')

# ═══════════════════════════════════════════════════════════════════════════════
add_h('10. Files Modified (This Session)', 2)
add_b('app/iww-hq/data/submissions-storage.ts \u2014 RealSubmission type extended with 18 '
      'new fields + 2 new aliases (KeyFeature, Award); mapRow reads new columns.')
add_b('app/api/submissions/route.ts \u2014 POST handler accepts and INSERTs all 18 new '
      'fields with arrJson() helper for empty-array \u2192 NULL coercion.')
add_b('app/listing/ListingDetailPage.tsx \u2014 imported KeyFeature + Award types; '
      'mapServerRow extended with 18 new fields; view object extended with 18 \u201creal*\u201d '
      'fields with sample fallback; render code wired at 7 places (header tags, pros/cons, '
      'sidebar starting price + free flags, support + training options, screenshot carousel, '
      'key features, all-features matrix).')
add_b('app/dashboard/new/NewListingClient.tsx \u2014 import swapped from ListingFormV2 to '
      'DashboardListingForm; plan-picker bullets refreshed with new PlanCaps flag names '
      '(hasKeyFeatures, hasPricingTiers, hasComplianceAndAwards).')
add_b('app/business/GetListedLanding.tsx \u2014 PlanKey type import re-pointed to '
      '../dashboard/new/form/DashboardListingForm.')
add_b('app/globals.css \u2014 removed @import \u201c./styles/listing-v2.css\u201d; added '
      '@import \u201c./styles/dashboard-form.css\u201d.')
add_b('app/styles/dashboard-shell.css \u2014 added .tp-main:has(.df-wrap) > .tp-content '
      'rule + .tp-content > .nl--form rule for the form-only internal-scroll layout, plus '
      'mobile override under the 900px breakpoint.')

# ═══════════════════════════════════════════════════════════════════════════════
add_h('11. Files Deleted (This Session)', 2)
add_b('app/business/ListingFormV2.tsx')
add_b('app/business/ListingForm.tsx')
add_b('app/business/form/ \u2014 entire directory (16 files: 2 root + 8 components + 9 '
      'steps minus 1 = 18 files actually).')
add_b('app/styles/listing-v2.css')

# ═══════════════════════════════════════════════════════════════════════════════
add_h('12. Deployment Checklist', 2)
add_b('CRITICAL: run database/migration-listings-v3.sql in phpMyAdmin BEFORE the form goes '
      'live in production. Each ALTER TABLE statement runs one at a time. If any column '
      'already exists (\u201cDuplicate column name\u201d error) just skip that line. Until '
      'this is run, POST /api/submissions will fail with \u201cUnknown column\u201d errors '
      'because the API now references the new columns.')
add_b('Hard-refresh /dashboard/new in the browser after deploy to pick up the new CSS bundle.')
add_b('Smoke-test: pick a plan \u2192 click \u201cFill demo data\u201d \u2192 navigate the '
      '7 steps \u2192 submit. Check /iww-hq/submissions for the new row, approve it, then '
      'visit /company/<slug> to confirm the listing renders the demo data instead of '
      'falling back to the Mailchimp sample.')

# ═══════════════════════════════════════════════════════════════════════════════
add_h('13. Known Non-Blocking Issues', 2)
add_b('Industries served, use cases, target company sizes, languages, mobile app flags, '
      'compliance, and awards are captured in the DB but the listing page does not yet '
      'render them as their own visual sections. The render code falls back to the '
      'hardcoded INDUSTRY donut, USE_CASES diamond cluster, COMPANY_SIZE bars, etc. \u2014 '
      'so they LOOK identical to the sample. A follow-up session will replace those '
      'visual blocks with submitter-driven renders. The data is captured today so the '
      'render-side migration is unblocked whenever it ships.')
add_b('Demo button leaves listingTypeIds and tagIds empty because they depend on l3Id and '
      'load asynchronously. User picks them manually after the demo fill auto-selects the '
      'category cascade.')
add_b('No undo button on the demo fill \u2014 if you click it after typing real data, your '
      'data is overwritten. Future polish: a confirm dialog when form is non-empty.')

doc.save(r'F:\infoWebWorld\InfoWebWorld_Session_Report_2026-03-23.docx')
print('Session 34 entry appended successfully.')

"""
Generate a Word doc documenting the www.infowebworld.com → dashboard issue,
so we can pick it up in a later session without re-tracing everything.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT = r"F:\infoWebWorld\WWW_Domain_Dashboard_Issue_2026-04-21.docx"

ACCENT = RGBColor(0xE8, 0x55, 0x3D)  # InfoWebWorld coral
INK = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x5C, 0x5C, 0x5C)
CODEBG = "F4F2EF"


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = INK if level > 1 else ACCENT
        run.font.name = "Calibri"
    return h


def add_para(doc, text, bold=False, italic=False, color=INK, size=11, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = "Calibri"
    r.bold = bold
    r.italic = italic
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        r.font.size = Pt(11)
        r.font.color.rgb = INK
        r.font.name = "Calibri"


def add_code(doc, text):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), CODEBG)
    pPr.append(shd)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    r.font.color.rgb = INK


def add_kv_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    for i, (k, v) in enumerate(rows):
        kc = table.rows[i].cells[0]
        vc = table.rows[i].cells[1]
        kc.width = Inches(2.2)
        vc.width = Inches(4.3)
        set_cell_bg(kc, "FAF5F0")
        kp = kc.paragraphs[0]
        kr = kp.add_run(k)
        kr.bold = True
        kr.font.size = Pt(10)
        kr.font.name = "Calibri"
        kr.font.color.rgb = INK
        vp = vc.paragraphs[0]
        vr = vp.add_run(v)
        vr.font.size = Pt(10)
        vr.font.name = "Calibri"
        vr.font.color.rgb = INK


# ─────────────────────────────────────────────────────────────

doc = Document()
styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
tr = title.add_run("WWW → Bare Domain Issue on Dashboard")
tr.font.size = Pt(22)
tr.font.color.rgb = ACCENT
tr.bold = True
tr.font.name = "Calibri"

sub = doc.add_paragraph()
sr = sub.add_run("Investigation Notes · 2026-04-21 · Deferred — Not Yet Applied")
sr.font.size = Pt(11)
sr.font.color.rgb = MUTED
sr.italic = True
sr.font.name = "Calibri"

doc.add_paragraph()

# ── 1. What the user sees ──
add_heading(doc, "1. What the user sees", level=1)
add_para(doc,
    "When logged into the dashboard on infowebworld.com, clicking the sidebar logo "
    "OR the \"Browse Site\" button sends the browser to www.infowebworld.com instead "
    "of staying on the bare infowebworld.com host.")
add_para(doc,
    "This ONLY happens from inside the dashboard. All other pages (/, /business, "
    "/plans, /contact, /company, /categories, etc.) render on bare and stay on bare. "
    "www does not appear anywhere else.")

# ── 2. Code audit ──
add_heading(doc, "2. Code audit — what was checked", level=1)
add_para(doc,
    "A full grep + manual read of every URL-emitting code path. Every file listed "
    "below was inspected for hardcoded hostnames, dynamic host construction, and "
    "redirect behavior.")

add_heading(doc, "2.1  Files checked (all clean — no www)", level=2)
add_bullets(doc, [
    "app/components/Navbar.tsx — logo <Link href=\"/\"> relative, no host",
    "app/[country]/dashboard/DashboardShell.tsx — sidebar logo <Link href={`/${country}`}>, browse-site <Link href={`/${country}`}>, logout router.push(`/${country}`) — all relative",
    "app/[country]/dashboard/layout.tsx — auth gate + redirect to /{country}/business on no session, relative path only",
    "app/layout.tsx — metadataBase = https://infowebworld.com (bare), openGraph.url = https://infowebworld.com (bare), no <base> tag, no service worker",
    "lib/oauth.ts — getCallbackUrl() reads PUBLIC_SITE_URL env or falls back to request.url.origin",
    "app/api/auth/[provider]/start/route.ts — redirect_uri built from getCallbackUrl()",
    "app/api/auth/[provider]/callback/route.ts — 302 to safeNext (relative) OR popup postMessage with relative next",
    "app/components/auth/SignupModal.tsx — window.location.href = data.next || nextUrl || '/dashboard' (all relative)",
    "app/components/auth/GoogleOneTap.tsx — POSTs /api/auth/google/one-tap then window.location.reload() (same host)",
    "middleware.ts — no host manipulation, only country + noindex logic",
    "next.config.ts — no www redirects or rewrites",
    "vercel.json — no redirect rules",
    "public/robots.txt — Sitemap: https://infowebworld.com (bare)",
    "app/sitemap.xml/route.ts + sitemap-pages.xml/route.ts — BASE = 'https://infowebworld.com' (bare)",
    "lib/user-auth.ts — userCookieHeader emits cookie with no Domain attribute → host-only",
])

add_heading(doc, "2.2  Literal string search", level=2)
add_code(doc,
    "grep -r \"www.infowebworld.com\" coming-soon/\n"
    "→ 0 matches (zero occurrences in any source file)")

add_heading(doc, "2.3  Env vars verified", level=2)
add_kv_table(doc, [
    ("PUBLIC_SITE_URL (.env.local)", "NOT SET"),
    ("PUBLIC_SITE_URL (Vercel prod)", "NOT SET (confirmed via `vercel env ls production`)"),
    ("Google OAuth redirect URI", "https://infowebworld.com/api/auth/google/callback (bare only — user confirmed)"),
    ("Session cookie scope", "Host-only on bare (no Domain attribute in Set-Cookie)"),
])

# ── 3. Conclusion ──
add_heading(doc, "3. Conclusion — where the www comes from", level=1)
add_para(doc,
    "The React/Next.js code does NOT emit www.infowebworld.com anywhere, and the "
    "OAuth flow is locked to bare (Google would reject a www redirect_uri since "
    "only bare is registered). Therefore the www appearance is not introduced by "
    "application code.")
add_para(doc, "The www must be coming from OUTSIDE the React code. Candidates:", bold=True)
add_bullets(doc, [
    "Chrome's address-bar autocomplete preferring www. for any site where www has ever responded (very aggressive; sticks per profile until cleared)",
    "A cached 301/308 from a previous test/config stored permanently in the browser",
    "Vercel domain configuration where both infowebworld.com and www.infowebworld.com are attached without a canonical redirect set between them",
    "An external link (email, LinkedIn, Product Hunt, social share) with www. that a visitor once clicked",
    "DNS — a www A/CNAME pointing to Vercel combined with no canonical preference",
])
add_para(doc,
    "The reason it is only noticed on the dashboard is circumstantial: the "
    "dashboard is the one surface reached exclusively after sign-in, so the "
    "browser may pick up www there via autocomplete/redirect cache while the "
    "user never notices it on the heavily-trafficked public pages. The Next.js "
    "relative links then keep the browser on whatever host it already landed on.",
    italic=True, color=MUTED)

# ── 4. Proposed fix ──
add_heading(doc, "4. Proposed fix (NOT yet applied)", level=1)
add_para(doc,
    "Add a www→bare 308 permanent redirect at the very top of middleware.ts. "
    "This catches ANY request on www.infowebworld.com at the edge and rewrites "
    "it to bare before page rendering — independent of Vercel domain config, "
    "DNS, browser autocomplete, or external backlinks.")

add_heading(doc, "4.1  Exact diff to apply", level=2)
add_para(doc, "File: middleware.ts — insert after line 27 (top of middleware function):", italic=True)
add_code(doc,
    "export function middleware(request: NextRequest) {\n"
    "  const host = request.headers.get('host') ?? ''\n"
    "\n"
    "  // Force bare domain — strip www. if the browser landed on www.infowebworld.com\n"
    "  if (host.startsWith('www.')) {\n"
    "    const url = request.nextUrl.clone()\n"
    "    url.host = host.slice(4)  // drop \"www.\"\n"
    "    return NextResponse.redirect(url, 308)\n"
    "  }\n"
    "\n"
    "  const { pathname } = request.nextUrl\n"
    "  // ... rest of existing middleware unchanged\n"
    "}")

# ── 5. Consequences ──
add_heading(doc, "5. Consequences of applying the fix", level=1)

add_heading(doc, "5.1  Positives", level=2)
add_bullets(doc, [
    "Any www.infowebworld.com/* visit → instant 308 redirect to bare. Issue fully gone — OAuth, logo click, browse-site, bookmarks, everything.",
    "Chrome stops suggesting www. in the address bar after 1-2 visits (308 is browser-cached).",
    "SEO gains a clean canonical hostname signal — no duplicate-content split.",
    "OAuth flow becomes more robust — any edge case landing on www gets rewritten before Google sees it.",
    "Analytics hostname dimension converges to a single value.",
])

add_heading(doc, "5.2  Things to be aware of", level=2)
add_bullets(doc, [
    "308 is permanent — browsers cache it HARD. If we ever want www to be primary in the future, reversal requires opposite redirect AND waiting out browser caches. Alternative: use 307 (temporary, not cached) if we want flexibility to reverse.",
    "One extra HTTP hop for www visitors (~50ms). Negligible.",
    "API routes on www are NOT caught by this middleware because the matcher excludes /api/*. This is actually correct: Google only sends OAuth callbacks to bare since that's the registered URI, so no www /api/* traffic is expected.",
    "No impact on Vercel preview deployments (*.vercel.app doesn't start with www.).",
    "Session cookies are already host-only on bare — no cookie-scope issues.",
    "If www.infowebworld.com is not currently routed to Vercel at all, this middleware is a harmless no-op — the problem would then be something else (browser autocomplete alone can't explain persistence).",
])

# ── 6. Decision needed ──
add_heading(doc, "6. Decision needed on next session", level=1)
add_kv_table(doc, [
    ("Option A", "Apply 308 permanent redirect (recommended — bare has been twice confirmed as canonical)"),
    ("Option B", "Apply 307 temporary redirect (same effect now, easier to reverse later)"),
    ("Option C", "Do nothing — live with www appearing on dashboard"),
])

add_para(doc, "")
add_para(doc,
    "Also pending on next session: investigate the continuous Singapore traffic "
    "in GA4 (likely Puppeteer/headless Chrome bots from AWS/GCP Singapore data "
    "centers — see separate note).",
    italic=True, color=MUTED)

# ── Footer ──
doc.add_paragraph()
foot = doc.add_paragraph()
fr = foot.add_run("— End of notes. Resume from section 6 when ready. —")
fr.font.size = Pt(9)
fr.font.color.rgb = MUTED
fr.italic = True
foot.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(OUT)
print(f"Saved: {OUT}")

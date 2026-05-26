"""
Generate a comprehensive Word doc describing the InfoWebWorld scraper bot
end-to-end: architecture, pipeline, cache, UI, screenshots, costs, every
issue hit, every fix applied, every commit. Lets us (Aadil + Claude)
pick up tomorrow with full context.

Run:
    python scripts/write-scraper-doc.py

Output:
    F:/infoWebWorld/InfoWebWorld_Scraper_Bot_Complete_Context.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(r"F:/infoWebWorld/InfoWebWorld_Scraper_Bot_Complete_Context.docx")

doc = Document()

# ───────────────────────────────────────────────────────────────────
# Style setup — body text Calibri 11, headings bold
# ───────────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def H1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x11, 0x18, 0x1C)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)


def H2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(0xE8, 0x55, 0x3D)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)


def H3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)


def P(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def CODE(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)


def BULLET(text):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(2)


def TABLE(rows, header=True):
    """Build a simple table. rows[0] = headers."""
    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
    tbl.style = "Light Grid Accent 1"
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = tbl.rows[r_idx].cells[c_idx]
            cell.text = str(cell_text)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    if r_idx == 0 and header:
                        run.bold = True
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)


# ═══════════════════════════════════════════════════════════════════
# COVER
# ═══════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
run = title.add_run("InfoWebWorld Scraper Bot")
run.bold = True
run.font.size = Pt(32)
run.font.color.rgb = RGBColor(0x11, 0x18, 0x1C)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

sub = doc.add_paragraph()
run = sub.add_run("Complete Build Context — for tomorrow's session")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

meta = doc.add_paragraph()
run = meta.add_run("Developer: Aadil Parmar  ·  AI pair: Claude Opus 4.7 (1M context)\n")
run.font.size = Pt(10)
run = meta.add_run("Repo: github.com/kakiso-official/infowebworld  ·  Branch: main\n")
run.font.size = Pt(10)
run = meta.add_run("Working dir: F:/infoWebWorld/launching/coming-soon\n")
run.font.size = Pt(10)
run = meta.add_run("Production: https://infowebworld.com  (Vercel)\n")
run.font.size = Pt(10)
run = meta.add_run("Database: cdbrisgy_infowebworld (cPanel MySQL)\n")
run.font.size = Pt(10)


# ═══════════════════════════════════════════════════════════════════
# 1. THE GOAL
# ═══════════════════════════════════════════════════════════════════
H1("1. What we're building and why")

P(
    "InfoWebWorld is a B2B directory across six L1 sectors — AI & ML, "
    "Software & SaaS, IT Services & Agencies, Local Businesses, Professional "
    "Services, and Startups. Each listing lives at /company/<slug> as a "
    "statically-rendered page with overview, key features, pricing tiers, "
    "FAQs, integrations, screenshots, and pros/cons. We had 94 AI/ML "
    "listings filled in by hand (a Python SQL generator with everything "
    "baked in as constants — fast to ship 26, painful to ship 100, "
    "impossible at 5,000)."
)

P(
    "The plan now is to populate the directory at scale — target 1,000 real "
    "companies per L1 (excluding Local Businesses, which is a different shape "
    "of listing). Each listing needs the same rich JSON we produced manually "
    "for the AI/ML 100 — but extracted from the company's own website by a "
    "scraping agent, not hand-coded. The agent uses Gemini 2.5 Pro for "
    "structured extraction with strict source citations so every claim is "
    "verifiable. Output is a per-listing UPDATE / INSERT into submissions "
    "which goes live the moment Vercel rebuilds the static page."
)

H2("Non-negotiables")
BULLET("Real info only — every claim must be quotable from the scraped page; null otherwise")
BULLET("SEO + AEO + GEO friendly — 200-350 word multi-paragraph descriptions, FAQs in question-form, schema-friendly")
BULLET("Hierarchical category mapping — L1 → L3 → L4 → L5, same depth as the 100 AI/ML mapping")
BULLET("Per-section re-scrape — refreshing one section must cost cents, not the full ~$0.05")
BULLET("Production-safe screenshots — uploaded to cPanel, not stored locally")
BULLET("Budget — Aadil has $250 of Gemini credit; pipeline must stay under it across all 5000 listings")


# ═══════════════════════════════════════════════════════════════════
# 2. ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════
H1("2. System architecture")

P(
    "Six pieces, decoupled. The worker runs LOCALLY on Aadil's machine "
    "(Playwright + Gemini API calls + uploads to cPanel). The admin UI runs "
    "in the Next.js dev server (also locally) and reads/writes the same DB. "
    "Vercel deploys serve the public /company pages — they read from the "
    "same DB at deploy time."
)

TABLE([
    ["#", "Component", "Purpose", "Where it runs"],
    ["1", "Database tables (5)", "Job queue + per-attempt sessions + step-by-step logs", "cPanel MySQL"],
    ["2", "Worker (scripts/scrape-worker.mjs)", "Long-running poller; runs the 13-step pipeline per queued job", "Local terminal"],
    ["3", "Pipeline (scripts/lib/scrape-pipeline.mjs)", "Crawl → screenshot → 5 extract passes → critique → validate", "Inside worker"],
    ["4", "File-based cache (.scrape-cache/)", "Per-job per-step JSON; makes retries free", "Local filesystem"],
    ["5", "API routes (/api/admin/scrape/*)", "CRUD + apply + section-specific requeue + deploy trigger", "Next.js dev server"],
    ["6", "Admin UI (/iww-hq/scraper)", "Two-pane Linear-style queue + live Gantt timeline", "Browser, dev server"],
])

H2("Why local-only for now")
P(
    "Vercel functions cap at 60 seconds. The full pipeline takes 30-60s per "
    "listing and Playwright Chromium is too big for Vercel's bundle. So the "
    "worker stays on Aadil's machine. The admin UI is dev-mode only for the "
    "same reason — it depends on the worker process. Upgrade path: deploy "
    "worker to a $5/mo DigitalOcean droplet later. Same code, no rewrite."
)


# ═══════════════════════════════════════════════════════════════════
# 3. DATABASE
# ═══════════════════════════════════════════════════════════════════
H1("3. Database — 5 new tables")

P(
    "Migration file: database/migration-scraper-tables.sql. "
    "All tables idempotent (CREATE TABLE IF NOT EXISTS). Already applied "
    "to cdbrisgy_infowebworld via phpMyAdmin."
)

TABLE([
    ["Table", "Purpose", "Cardinality"],
    ["scrape_jobs", "One row per company in the system. Status, cached extracted JSON, screenshot URLs.", "1 per company (~5,000 target)"],
    ["scrape_sessions", "One row per scrape attempt. Multiple sessions per job over time.", "1-10 per job"],
    ["scrape_session_steps", "One row per step within a session. Drives the Gantt timeline.", "~14 per session"],
    ["companies_seed", "Pre-discovered candidate pool (NOT YET POPULATED).", "Target 5,000 across L1s"],
    ["discovery_sessions", "Each batch of company-discovery (YC, Product Hunt, etc.).", "Few per L1"],
])

H2("Key columns and shapes")
H3("scrape_jobs")
CODE(
    "id, slug (unique), company_name, website, category_l1, category_slug, "
    "status (queued|running|review|applied|failed|skipped), last_session_id, "
    "total_sessions, total_cost_usd,\n"
    "current_extracted_json (the snapshot of the latest applied extraction),\n"
    "current_source_citations,\n"
    "current_screenshot_home_url, current_screenshot_secondary_url,\n"
    "queued_at, applied_at"
)
H3("scrape_sessions")
CODE(
    "id, scrape_job_id, status (running|success|failed|cancelled|review),\n"
    "total_steps, failed_steps, total_input_tokens, total_output_tokens,\n"
    "total_cost_usd, pages_fetched,\n"
    "extracted_json, source_citations, critique_flags,\n"
    "screenshot_home_url, screenshot_secondary_url,\n"
    "llm_provider, model_version,\n"
    "started_at, finished_at, duration_ms, error_summary,\n"
    "applied_at, applied_by"
)
H3("scrape_session_steps")
CODE(
    "id, session_id, step_name, step_index, step_type (crawl|extract|screenshot|critique|validate|save),\n"
    "status (running|ok|error|skipped),\n"
    "input_url, input_prompt_preview, output_status_code, output_bytes,\n"
    "output_excerpt, input_tokens, output_tokens, cost_usd, error_message,\n"
    "started_at, finished_at, duration_ms"
)


# ═══════════════════════════════════════════════════════════════════
# 4. THE 13-STEP PIPELINE
# ═══════════════════════════════════════════════════════════════════
H1("4. Pipeline — 13 steps per listing")

P(
    "Every scrape session runs through these steps in order. Each step gets "
    "its own row in scrape_session_steps. Crawl and extract steps are CACHED "
    "to disk; retries skip cached steps for $0."
)

TABLE([
    ["#", "Step", "Type", "Cost (Pro)", "What it does"],
    ["1",  "crawl-home",         "crawl",      "$0",      "Playwright fetches the homepage."],
    ["2",  "crawl-about",        "crawl",      "$0",      "Tries /about, /about-us, /company, /team, /who-we-are."],
    ["3",  "crawl-pricing",      "crawl",      "$0",      "Tries /pricing, /plans, /buy, /price, etc."],
    ["4",  "crawl-features",     "crawl",      "$0",      "Tries /features, /product, /platform, /capabilities, etc."],
    ["5",  "crawl-faq",          "crawl",      "$0",      "Tries /faq, /faqs, /help, /support, etc."],
    ["6",  "crawl-integrations", "crawl",      "$0",      "Tries /integrations, /apps, /marketplace, /partners, etc."],
    ["7",  "screenshot-home",    "screenshot", "$0",      "Captures 1440x900 JPG; uploads to cPanel via /api/upload."],
    ["8",  "screenshot-secondary","screenshot","$0",      "Captures /pricing or /features (whichever exists)."],
    ["9",  "extract-base",       "extract",    "~$0.005", "Tagline + 200-350 word description + founded + HQ + team + socials."],
    ["10", "extract-pricing",    "extract",    "~$0.008", "pricing_model + pricing_tiers + starting_price + has_free_trial/version."],
    ["11", "extract-features",   "extract",    "~$0.012", "10-15 key_features {name,description} + 15-25 features (strings) + integrations + languages + compliance + apps."],
    ["12", "extract-faqs",       "extract",    "~$0.006", "8 FAQs as {question, answer}, mix of pricing/capability/comparison/logistics."],
    ["13", "extract-classify",   "extract",    "~$0.004", "header_tags + industries_served + use_cases + target_company_sizes + pros + cons + support_channels + training_options."],
    ["14", "self-critique",      "critique",   "~$0.008", "Gemini audits the extracted JSON for hallucinated claims."],
    ["15", "validate-citations", "validate",   "$0",      "Client-side check: every non-null field has _source_url + _source_quote."],
    ["16", "save-session",       "save",       "$0",      "Writes final JSON + citations + screenshot URLs to scrape_sessions."],
])

P(
    "Total per-listing cost: ~$0.04–0.06 with Gemini 2.5 Pro. With cache: a "
    "retry that touches only one section is ~$0.005–0.012. A screenshots-only "
    "refresh is $0 of LLM (just Playwright + upload)."
)


# ═══════════════════════════════════════════════════════════════════
# 5. THE REAL-INFO GUARANTEE
# ═══════════════════════════════════════════════════════════════════
H1("5. Real-info guarantee — three layers")

P(
    "This is THE non-negotiable. Aadil's hard rule from day one was no "
    "gimmicks, no hallucinated SOC 2 / HIPAA / integrations. The pipeline "
    "enforces it at three independent layers so a failure in any one is "
    "caught by the next."
)

H3("Layer A — Schema-required source citations")
P(
    "Every Gemini extract pass uses responseSchema (Google's structured "
    "output) with required fields _source_url and _source_quote per "
    "non-null value. If Gemini can't quote it, the value must be null. "
    "Schema validation refuses the JSON otherwise."
)

H3("Layer B — Strict prompts")
P(
    "Every prompt has a SYSTEM_EXTRACT preface that says: \"REAL INFO ONLY. "
    "If a fact is not present in the supplied HTML, set value to null. Do "
    "NOT infer pricing, compliance certifications, integrations, or "
    "features from prior knowledge of the company. _source_quote must be "
    "a VERBATIM substring of the cleaned HTML, max 200 chars.\""
)

H3("Layer C — Self-critique pass")
P(
    "After all 5 extract passes, a 6th pass (self-critique) gets the "
    "assembled JSON back and is asked to flag any claim that looks "
    "hallucinated — compliance certs without evidence, suspicious generic "
    "pricing, boilerplate phrasings. Output: list of flagged fields + an "
    "overall_quality score. If anything is flagged, the session status is "
    "set to 'review' (not 'success') so it surfaces in the UI for manual "
    "check before Apply."
)


# ═══════════════════════════════════════════════════════════════════
# 6. CACHE SYSTEM
# ═══════════════════════════════════════════════════════════════════
H1("6. Cache — what makes retries cheap")

P(
    "Initial cost concern: re-scraping a listing because one section came "
    "back weak shouldn't cost the full ~$0.05. Solution: file-based per-job "
    "per-step cache at .scrape-cache/<jobId>/<step>.json with 24h TTL. "
    "Stores the step's _data output."
)

H3("How it works")
BULLET("Before each cacheable step, check getCached(jobId, stepName). If hit, log '[cache] step-name' and return immediately ($0).")
BULLET("On cache hit, the pipeline still inserts a step row marked status='ok', output_excerpt='cached (from previous session)', cost_usd=0 — so the Gantt timeline still shows the step.")
BULLET("On cache miss, run the step normally and setCached(jobId, stepName, result._data) after success.")

H3("Cacheable steps")
P(
    "All crawl steps (1-6), all extract steps (9-13), and both screenshot "
    "steps (7-8). The screenshot cache value is the public cPanel URL. "
    "Self-critique is NOT cached (depends on the freshly assembled JSON). "
    "Validate and save are NOT cached (bookkeeping)."
)

H3("Cache invalidation — per-section requeue")
P(
    "The requeue endpoint accepts a 'sections' array. Each section maps to "
    "a list of cache files to delete:"
)
CODE(
    "all          → wipe the entire job cache (full fresh ~$0.05)\n"
    "screenshots  → screenshot-home + screenshot-secondary ($0 LLM)\n"
    "base         → crawl-home + crawl-about + extract-base (~$0.005)\n"
    "pricing      → crawl-pricing + extract-pricing (~$0.008)\n"
    "features     → crawl-features + crawl-integrations + extract-features (~$0.012)\n"
    "integrations → crawl-integrations + extract-features (~$0.012)\n"
    "faqs         → crawl-faq + extract-faqs (~$0.006)\n"
    "classify     → extract-classify (~$0.004)"
)


# ═══════════════════════════════════════════════════════════════════
# 7. SCREENSHOTS
# ═══════════════════════════════════════════════════════════════════
H1("7. Screenshots — capture, upload, serve")

P(
    "The 100 AI/ML listings already used the existing /api/upload endpoint, "
    "which proxies to cPanel storage via IP+Host header (because Vercel "
    "rewrites infowebworld.com → its own edge). Returns a /api/file/<path> "
    "URL that the Next.js /api/file proxy can re-fetch from cPanel and "
    "edge-cache. The new scraper uses the SAME endpoint."
)

H3("Capture")
BULLET("Playwright headless Chromium, 1440x900, dismisses cookie banners, waits for network-idle + 1.5s settle.")
BULLET("Writes JPG (quality 82) to public/scrape-screenshots/<slug>-<kind>.jpg as a local backup/debug artefact.")
BULLET("kind = 'home' or 'secondary'.")

H3("Upload")
BULLET("Worker reads the local JPG, builds multipart/form-data, POSTs to ${SITE_BASE}/api/upload.")
BULLET("Default SITE_BASE = http://localhost:3000 (Aadil's dev server). Override with SITE_BASE=https://infowebworld.com if needed.")
BULLET("Endpoint validates type+size, forwards via http (not fetch — fetch strips the Host header) to CPANEL_IP=104.152.168.46 with Host=infowebworld.com.")
BULLET("Returns /api/file/logos/<id>_<timestamp>.jpg style URL.")

H3("Display")
BULLET("Apply route writes the returned URL into submissions.screenshots (JSON array).")
BULLET("Listing page reads r.screenshots, parses, shows in the carousel.")
BULLET("Same URL works on localhost and infowebworld.com because /api/file proxies through Next.js on both.")

H3("Scraper UI badges")
BULLET("Each screenshot thumbnail in /iww-hq/scraper has a corner badge:")
BULLET("  GREEN 'on cPanel — prod-safe' when URL starts with /api/file/")
BULLET("  RED 'LOCAL ONLY — won't load on infowebworld.com' when URL starts with /scrape-screenshots/")
BULLET("If red shows up: dev server wasn't running when worker tried to upload. Pick 'screenshots only' from section dropdown, re-scrape, check worker terminal for the '✓ uploaded' line.")


# ═══════════════════════════════════════════════════════════════════
# 8. ADMIN UI
# ═══════════════════════════════════════════════════════════════════
H1("8. Admin UI — /iww-hq/scraper")

P(
    "Two-pane Linear-style layout. Top bar with status counters + worker "
    "status + 7-day spend + Deploy + Add Companies. Filter strip with L1 "
    "+ status + search. Body: queue list on the left, listing detail on the "
    "right with sessions accordion + Gantt timeline."
)

H3("Top bar pieces")
BULLET("Status pill counters (queued / running / review / applied / failed / skipped) — clickable filters")
BULLET("Worker online/offline dot — green when last activity < 5 min ago, red otherwise. Shows 'npm run scrape:worker' hint if offline.")
BULLET("7d spend meter — sum of total_cost_usd across all sessions in last 7 days")
BULLET("Deploy to Vercel button — fires the existing /api/admin/deploy endpoint (POST to VERCEL_DEPLOY_HOOK_URL). Required after Apply to refresh /company/<slug> on infowebworld.com.")
BULLET("+ Add Companies button — opens modal for bulk CSV-style entry")

H3("Listing detail (right pane)")
BULLET("Hero with logo + name + website + category chips + status pill")
BULLET("Action bar: 'Scrape now' (cheapest retry) + chevron dropdown (per-section) + 'Apply latest to DB'")
BULLET("Screenshots: home + secondary side-by-side with prod-safety badges")
BULLET("Sessions accordion: newest first, click to expand the Gantt timeline")
BULLET("Each step in the timeline: status icon, name, horizontal bar (length = duration), tokens, cost — click to drill into prompt/output/error")
BULLET("Latest extraction JSON in a collapsible <details> for raw inspection")


# ═══════════════════════════════════════════════════════════════════
# 9. CLI SCRIPTS
# ═══════════════════════════════════════════════════════════════════
H1("9. CLI scripts — terminal usage")

H3("Setup (one-time)")
CODE(
    "cd F:/infoWebWorld/launching/coming-soon\n"
    "npx playwright install chromium    # downloads the headless browser"
)

H3("Bootstrap — import the 94 enriched AI/ML listings as 'applied'")
CODE(
    "npm run scrape:seed-applied -- --dry-run    # preview\n"
    "npm run scrape:seed-applied                  # actually do it"
)

H3("Single-shot scrape (CLI)")
CODE(
    "npm run scrape:listing -- --slug=loom\n"
    "npm run scrape:listing -- --slug=loom --model=gemini-2.5-flash\n"
    "npm run scrape:listing -- --slug=loom --dry-run\n"
    "# auto-create job if not yet in scrape_jobs:\n"
    "npm run scrape:listing -- --slug=loom --website=https://www.loom.com --category-l1=software-and-saas"
)

H3("Worker (the main loop)")
CODE(
    "npm run scrape:worker\n"
    "# Options:\n"
    "#   --concurrency=2          (default 1, max 5)\n"
    "#   --model=gemini-2.5-flash  (default gemini-2.5-pro)\n"
    "#   --poll=8000               (default 10000ms)\n"
    "#   --daily-cap=20            ($ cap per UTC day, default $20)\n"
    "#   --job-cap=0.50            ($ cap per session, default $0.50)\n"
    "#   --l1=software-and-saas    (filter to one L1)"
)

P(
    "Worker is a long-running Node process. Code changes do NOT hot-reload. "
    "After every git push (or local code change), Ctrl-C the worker terminal "
    "and re-run `npm run scrape:worker`."
)


# ═══════════════════════════════════════════════════════════════════
# 10. APPLY — THE LIVE LISTING
# ═══════════════════════════════════════════════════════════════════
H1("10. Apply — writing to live submissions")

P(
    "Apply (button on the listing detail pane, POST /api/admin/scrape/jobs/[id]/apply) "
    "is what makes the data go LIVE on the directory."
)

H3("Flow")
BULLET("Reads the latest session's extracted_json + source_citations + screenshot URLs.")
BULLET("Looks up the matching submissions row by slug.")
BULLET("If submissions row EXISTS → UPDATE the rich fields (tagline, description, pricing_tiers, faqs, integrations, screenshots, etc.)")
BULLET("If NO submissions row → INSERT a new one with defaults: status='active', payment_status='completed', listing_mode='product', plan_id = cheapest active plan, category_id resolved from category_slug → L1 fallback, country_id from extracted hq_country_code → US fallback, contact_name=company_name, email=contact@<domain>.")
BULLET("Calls revalidatePath('/company/<slug>') so Vercel rebuilds the static page. (Caveat: only takes effect when Apply hits Vercel; local apply requires the Deploy button to refresh production.)")

H3("Null handling")
P(
    "Fields the scraper couldn't verify come back null. The apply route "
    "SKIPS null values (uses DB defaults / preserves existing values) instead "
    "of writing NULL. This avoids 'Column X cannot be null' errors on the "
    "TINYINT(1) booleans (has_ios_app, has_android_app, has_free_trial, "
    "has_free_version) which are NOT NULL DEFAULT 0."
)

H3("Boolean coercion")
P(
    "Booleans are explicitly coerced to 0/1 instead of true/false to avoid "
    "mysql2 driver edge cases. JSON fields are JSON.stringify'd. Everything "
    "else passes through as-is."
)

H3("key_features shape coercion (important)")
P(
    "Listing page expects keyFeatures = [{name, description}]. Older "
    "extractions (and the AI/ML 100 hand-written enrichments) stored plain "
    "string arrays. Apply now coerces strings to {name, description: ''} "
    "before writing, so legacy data and new scrapes both render correctly."
)


# ═══════════════════════════════════════════════════════════════════
# 11. SECTION DROPDOWN
# ═══════════════════════════════════════════════════════════════════
H1("11. Section-specific re-scraping")

P(
    "Big cost-saver. Next to 'Scrape now' is a chevron dropdown listing "
    "every section the user can refresh independently."
)

TABLE([
    ["Option", "Cost", "Steps cleared from cache"],
    ["all sections (full fresh)",       "~$0.05",  "Everything in the job cache"],
    ["screenshots only",                "$0 LLM",  "screenshot-home + screenshot-secondary"],
    ["description (base)",              "~$0.005", "crawl-home + crawl-about + extract-base"],
    ["pricing",                         "~$0.008", "crawl-pricing + extract-pricing"],
    ["features + integrations",         "~$0.012", "crawl-features + crawl-integrations + extract-features"],
    ["integrations only",               "~$0.012", "crawl-integrations + extract-features"],
    ["FAQs",                            "~$0.006", "crawl-faq + extract-faqs"],
    ["classification",                  "~$0.004", "extract-classify"],
])

P(
    "Implementation: requeue endpoint takes { sections: string[] } body. "
    "Each section name maps to an array of cache file names to delete. After "
    "clearing, the job status is set to 'queued' and the worker picks it up. "
    "Pipeline runs all steps but most are cache hits — only the cleared ones "
    "actually fetch / extract."
)


# ═══════════════════════════════════════════════════════════════════
# 12. ISSUES HIT + FIXED
# ═══════════════════════════════════════════════════════════════════
H1("12. Every issue we hit and how we fixed it")

issues = [
    ("Migration paste-error — SQL inline comments swallowed columns",
     "phpMyAdmin flattened newlines on paste; the inline '-- ...' end-of-line comments ate everything after them on each statement, breaking the CREATE TABLE.",
     "Rewrote migration-scraper-tables.sql with all comments on their own lines (block style above each table). Idempotent CREATE TABLE IF NOT EXISTS so re-running is safe."),
    ("Gemini 2.5 Pro: 'Budget 0 is invalid'",
     "I defaulted thinkingBudget=0 in the Gemini wrapper to disable hidden reasoning (saves tokens). Works on Flash. Pro REQUIRES thinking mode — returns 400 INVALID_ARGUMENT.",
     "Default thinkingBudget=null. Only emit thinkingConfig if explicitly set. Silently drop thinkingBudget=0 for Pro models."),
    ("Failed sessions burnt full $ on retry",
     "Initial implementation re-did everything from scratch on every scrape attempt. A single failed step = $0.04 lost.",
     "Built file-based per-job per-step cache at .scrape-cache/. Crawl + extract steps cached. Retries skip cached steps for $0. Section dropdown clears only specific cache files."),
    ("Apply failed for new listings — 'No submissions row for slug=X'",
     "Apply assumed the submissions row already existed. Loom didn't have one (only the 94 AI/ML had hand-written rows).",
     "Apply now auto-INSERTs a new submissions row with sensible defaults when no existing row found. Status=active, payment=completed, plan=cheapest, category resolved from L3 slug → L1 fallback."),
    ("Column 'has_ios_app' cannot be null",
     "Scraper returns null when it can't verify a boolean. But submissions.has_ios_app is NOT NULL DEFAULT 0.",
     "Apply route skips undefined AND null values entirely (DB defaults win). Booleans explicitly coerced to 0/1."),
    ("key_features rendered as empty rows",
     "Listing page expects KeyFeature = {name, description}[]. Scraper output string[]. The .map(kf => ({name: kf.name})) gave undefined for each name.",
     "Schema changed: key_features is now an array of {name, description, source_quote} objects. Separate plain-string 'features' array for the comprehensive feature list. Apply route coerces strings to {name, description: ''} for backward compat with the 26 hand-written enrichments."),
    ("Description too short / no paragraphs",
     "First scrapes returned ~80-word single-paragraph descriptions. Aadil wanted 200+ words with proper paragraphs.",
     "Prompt now demands 200-350 words in 3 paragraphs (what it is / what it does / credibility) separated by \\n\\n. Listing page render splits on /\\n{2,}/ into separate <p> tags."),
    ("Integrations always empty",
     "Pipeline only crawled /pricing /features /faq /about. Most SaaS sites keep integrations at /integrations or /apps which we never touched. Even with strict prompt rules, the extract had nothing to work from.",
     "Added INTEGRATIONS_FALLBACKS chain (/integrations, /apps, /marketplace, /partners, /works-with, /ecosystem, /connectors, /extensions). New crawl-integrations step. Feeds that HTML into the features extract. Integrations prompt strengthened: minimum 8, scan logo walls + img alt-text + feature copy."),
    ("Screenshots not showing on infowebworld.com",
     "Worker saved screenshots to public/scrape-screenshots/ on Aadil's machine. .gitignored — not in Vercel deploy. So /company/<slug> on prod points to URLs that 404.",
     "scrape-screenshot.mjs uploads via the existing /api/upload endpoint (same one capture-screenshots.mjs uses). Returns /api/file/<path> URL. Pipeline stores THAT in scrape_sessions. Apply writes it to submissions.screenshots. Works on both localhost and infowebworld.com."),
    ("Silent upload failures (dev server offline)",
     "If npm run dev wasn't running when worker called upload, ECONNREFUSED was caught and we fell back to the local path. No warning, listing page on prod 404'd quietly.",
     "Added explicit console.log lines: 'uploading <file> → <url>' before each upload, '✓ uploaded <file> → <url>' on success, '✗ upload FAILED for <file>: <reason>' on failure. New ScreenshotBadge in the Scraper UI: green 'on cPanel — prod-safe' or red 'LOCAL ONLY' so you can tell at a glance."),
    ("Worker not picking up code changes",
     "Worker is a long-running Node process. Edits to scrape-pipeline.mjs / scrape-gemini.mjs etc. don't hot-reload — the worker runs the code it started with.",
     "Always Ctrl-C and re-run `npm run scrape:worker` after a git push. Documented in the run instructions."),
    ("Vercel cached /company/<slug> doesn't refresh after Apply",
     "/company/<slug> is statically rendered at deploy time (revalidate=172800, dynamicParams=false). Apply's revalidatePath() runs on the LOCAL dev server — only refreshes localhost.",
     "Added Deploy button in the Scraper top bar. Fires existing /api/admin/deploy which POSTs to VERCEL_DEPLOY_HOOK_URL. Triggers Vercel build which regenerates all /company pages from current DB. ~1-2 min then prod is fresh."),
    ("Screenshots-only re-scrape was missing from UI",
     "Section dropdown had pricing, features, faqs, classify but no way to refresh screenshots without re-running everything.",
     "Added 'screenshots only — $0 LLM' option. Also made screenshot steps use cachedStep (cached value is the cPanel URL) so non-screenshot section re-scrapes don't redo Playwright + upload."),
]

for title, problem, fix in issues:
    H3(title)
    P("Problem: " + problem)
    P("Fix: " + fix)


# ═══════════════════════════════════════════════════════════════════
# 13. CURRENT STATE
# ═══════════════════════════════════════════════════════════════════
H1("13. Where we are right now")

H2("Working end-to-end")
BULLET("Loom successfully scraped, applied, deployed. Screenshots on cPanel. Listing live at /company/loom.")
BULLET("94 AI/ML listings imported to scrape_jobs as 'applied' (won't be re-scraped).")
BULLET("Pipeline: 13 steps + critique + validate + save = ~$0.04-0.06 per Pro scrape.")
BULLET("Section dropdown live: screenshots only, base, pricing, features, integrations, faqs, classify, all.")
BULLET("Cache works: cached retries cost $0 for everything except critique (~$0.008).")
BULLET("Deploy button live: refreshes Vercel static pages on demand.")
BULLET("Per-screenshot prod-safety badges in admin UI.")

H2("NOT done yet — queued for next session(s)")
BULLET("Discovery Agent: 1,000-per-L1 candidate pool (companies_seed table is built but unpopulated). Sources: YC public directory, Product Hunt, Clutch, G2 categories, BetaList, Crunchbase. Each candidate gets hierarchical category mapping (L3→L4→L5) + responsiveness score (likelihood the company will respond to a 'claim your listing' email).")
BULLET("Per-field review grid: accept/reject/edit individual extracted fields before Apply, with source URL + quote inline.")
BULLET("Activity stream tab: live cross-job event feed.")
BULLET("Dashboards tab: per-L1 completion progress, cost over time, success rate, top failure reasons, categorisation confidence histogram.")
BULLET("Migration for the 26 hand-written AI/ML key_features (strings) → object shape for the listing page to render names properly.")
BULLET("Categorisation pass: when discovery adds a candidate, a Gemini Flash call maps it to L3/L4/L5 with confidence + reasoning, same depth as aiml-listing-mapping-notes.md.")
BULLET("Responsiveness score: compute from public signals (founded year, funding stage, team size, blog activity, G2/Capterra/PH presence). Sort candidates DESC so we scrape the most-likely-to-claim first.")

H2("Open known risks / TODOs")
BULLET("If a site changes between scrapes the 24h cache might return stale data. TTL is per-step — can be reduced if needed.")
BULLET("Self-critique sometimes over-flags valid claims (compliance fragments etc.). Currently sets session to 'review' — manual eyeball decides apply.")
BULLET("No prompt versioning in the cache. If we improve a prompt, we need to manually clear .scrape-cache/ to re-use it.")
BULLET("Cost cap per session is $0.30 (default). Rich content + critique + ~$0.04 base = comfortable margin. Bump to $1 if we ever increase token budgets.")


# ═══════════════════════════════════════════════════════════════════
# 14. FILE INVENTORY
# ═══════════════════════════════════════════════════════════════════
H1("14. Complete file inventory")

H2("Database")
BULLET("database/migration-scraper-tables.sql — 5 tables, idempotent")

H2("Library (shared by API + scripts)")
BULLET("lib/scrape/types.ts — TypeScript types matching the 5 DB tables, plus ExtractedListing / SourceCitation / GeminiModel")

H2("Scripts library")
BULLET("scripts/lib/scrape-env.mjs — loads .env.local into a plain object")
BULLET("scripts/lib/scrape-db.mjs — mysql2/promise wrapper, q / q1 / exec helpers")
BULLET("scripts/lib/scrape-gemini.mjs — Gemini REST client with structured-output support, cost calculator, retry on 429/5xx")
BULLET("scripts/lib/scrape-crawler.mjs — Playwright singleton + fetchPage + fetchWithFallback + PRICING_FALLBACKS, FEATURES_FALLBACKS, FAQ_FALLBACKS, ABOUT_FALLBACKS, INTEGRATIONS_FALLBACKS")
BULLET("scripts/lib/scrape-screenshot.mjs — captureScreenshot (Playwright) + uploadScreenshot (multipart POST to /api/upload)")
BULLET("scripts/lib/scrape-clean.mjs — cleanHtml (strips script/style/svg/data: URLs), slugifyName, canonicalUrl")
BULLET("scripts/lib/scrape-schemas.mjs — 6 Gemini responseSchemas (baseInfo, pricing, features, faqs, classification, critique) + categorySchema for discovery")
BULLET("scripts/lib/scrape-cache.mjs — file-based per-job per-step cache, 24h TTL, getCached/setCached/clearJobCache/summarizeJobCache")
BULLET("scripts/lib/scrape-pipeline.mjs — the orchestrator. 13 steps + critique + validate + save. The biggest single file in the system.")

H2("CLI entry points")
BULLET("scripts/scrape-listing.mjs — single-shot CLI by slug or job id")
BULLET("scripts/scrape-worker.mjs — long-running poller, concurrency/daily-cap/model/L1 filter options")
BULLET("scripts/seed-applied-jobs.mjs — bootstrap import of existing enriched submissions as 'applied'")
BULLET("scripts/write-scraper-doc.py — generates this Word doc")

H2("API routes")
BULLET("app/api/admin/scrape/jobs/route.ts — GET list (with l1/status/q/limit/offset filters) + POST bulk add")
BULLET("app/api/admin/scrape/jobs/[id]/route.ts — GET detail (job + sessions) + PATCH limited field updates")
BULLET("app/api/admin/scrape/jobs/[id]/requeue/route.ts — POST sections=[...] to clear specific cache files + queue")
BULLET("app/api/admin/scrape/jobs/[id]/apply/route.ts — POST writes the extraction into submissions; auto-INSERTs if missing; revalidatePath after success")
BULLET("app/api/admin/scrape/sessions/[id]/route.ts — GET session + steps (used by the Gantt timeline polling)")
BULLET("app/api/admin/scrape/stats/route.ts — GET aggregate status counters + 7d cost + worker-likely-online flag")

H2("Admin UI")
BULLET("app/iww-hq/scraper/page.tsx — server entry, dynamic=force-dynamic")
BULLET("app/iww-hq/scraper/ScraperShell.tsx — main client component, polling, two-pane layout, DeployButton, AddCompaniesModal trigger, SectionMenu, ScreenshotBadge")
BULLET("app/iww-hq/scraper/SessionTimeline.tsx — Gantt-style horizontal bar chart per step, live polling 1.5s while session.status='running'")
BULLET("app/iww-hq/scraper/AddCompaniesModal.tsx — CSV paste, parses to {slug, company_name, website, category_l1, category_slug}, bulk POSTs to /api/admin/scrape/jobs")

H2("Styles")
BULLET("app/styles/scraper.css — ~700 lines, .scrp-* namespace, no Tailwind, solid colors, 1.5px borders, no glow per design rules")

H2("Modified existing files")
BULLET("app/listing/ListingDetailPage.tsx — description render now splits on /\\n{2,}/ into multiple <p> tags")
BULLET("package.json — added 3 npm scripts: scrape:listing, scrape:worker, scrape:seed-applied")
BULLET(".gitignore — added /.scrape-cache/ and /public/scrape-screenshots/")


# ═══════════════════════════════════════════════════════════════════
# 15. COMMITS LOG
# ═══════════════════════════════════════════════════════════════════
H1("15. Commits pushed (in order)")

commits = [
    ("9d6da8a", "feat(scraper): monster — Gemini agent + admin UI + per-step cache + live timeline"),
    ("8f845c8", "feat(scraper): rich key_features, 200-word descriptions, deeper prompts"),
    ("0c5e609", "fix(scraper): dedicated /integrations crawl + stronger extract prompt"),
    ("e17e0ad", "feat(scraper): per-section re-scrape — stop burning $ on full retries"),
    ("5fbd1aa", "fix(scraper): upload screenshots to cPanel via /api/upload — production-safe"),
    ("63c8f73", "feat(scraper): screenshots-only section + cache screenshot URLs"),
    ("be36c7e", "fix(scraper): diagnose silent screenshot upload failures"),
    ("1eb0607", "feat(scraper): Deploy button in scraper top bar"),
]
for sha, msg in commits:
    P(f"{sha}  —  {msg}")


# ═══════════════════════════════════════════════════════════════════
# 16. HOW TO USE — STEP BY STEP
# ═══════════════════════════════════════════════════════════════════
H1("16. Daily workflow — Aadil's step-by-step")

H3("Setup (once)")
BULLET("npx playwright install chromium")
BULLET("Verify .env.local has DATABASE_*, GEMINI_API_KEY")
BULLET("Migration applied (already done via phpMyAdmin)")
BULLET("npm run scrape:seed-applied (already done — imports the 94 AI/ML)")

H3("Every session")
BULLET("Terminal #1: cd F:/infoWebWorld/launching/coming-soon && npm run dev")
BULLET("Terminal #2: cd F:/infoWebWorld/launching/coming-soon && npm run scrape:worker")
BULLET("Browser: http://localhost:3000/iww-hq/scraper")

H3("Adding companies")
BULLET("Click '+ Add Companies' in the top bar")
BULLET("Paste CSV-style lines: slug, company_name, website, category_l1, category_slug")
BULLET("One per line. Header line auto-skipped. Lines starting with # ignored. Duplicates by slug skipped silently.")
BULLET("Click 'Add to queue'")

H3("Scraping a listing")
BULLET("Click the row in the left list to select")
BULLET("Click 'Scrape now' (cheapest retry — uses all cache) OR ▾ dropdown for a specific section")
BULLET("Worker picks it up in ~10 seconds; watch the Gantt timeline fill in")
BULLET("Each step shows duration + tokens + cost as it runs")
BULLET("Session finishes with status 'success' or 'review' (review = self-critique flagged something)")

H3("Applying to live submissions")
BULLET("If status is 'review', expand the session and check the critique flags + extracted JSON")
BULLET("Once happy, click 'Apply latest to DB'")
BULLET("Apply writes the data to submissions, sets status='applied'")
BULLET("Screenshot URLs flow through to submissions.screenshots")

H3("Making changes visible on infowebworld.com")
BULLET("Click 'Deploy to Vercel' in the top bar (confirms first)")
BULLET("Vercel rebuilds the static /company pages from current DB")
BULLET("~1-2 minutes later, /company/<slug> on infowebworld.com shows the new data")

H3("After each git push (or local code change)")
BULLET("Ctrl-C the worker terminal")
BULLET("Restart with npm run scrape:worker — code changes don't hot-reload")


# ═══════════════════════════════════════════════════════════════════
# 17. NEXT SESSION PRIORITIES
# ═══════════════════════════════════════════════════════════════════
H1("17. Tomorrow's priorities — pick from these")

priorities = [
    ("Discovery Agent",
     "Highest leverage. Build scripts/discover.mjs that ingests a source (YC public directory, Product Hunt category page, Clutch top list, manual CSV) and populates companies_seed. Each candidate gets a hierarchical category mapping (Gemini Flash call with L3 children of the target L1, then L4 children of the picked L3, then L5) + a responsiveness score from public signals. Sort DESC by responsiveness. UI tab 'Candidates' filters by L1 + status + score, bulk 'Promote top N to scrape queue'."),
    ("Per-field review grid",
     "Aadil's 'real info only' check still relies on him reading the extracted JSON. A grid that shows every extracted field with the source URL + quote inline, with ✓ accept / ✗ reject / ✏ edit per field. Apply only writes the accepted fields. Rejected fields get cleared from cache so the next re-scrape doesn't reuse them."),
    ("AI/ML 26 backfill",
     "The 26 hand-written enrichments have key_features as plain string arrays. Migration: SELECT slug, key_features FROM submissions WHERE pricing_tiers IS NOT NULL — for each row, if JSON_TYPE(JSON_EXTRACT(key_features, '$[0]')) = 'STRING', wrap each string as {name: <string>, description: ''} and UPDATE. Single SQL or one-shot Node script."),
    ("Dashboards tab",
     "Cost over time line chart, success rate gauge, top failure reasons bar chart, per-L1 completion funnel, average tokens per pass. Polling endpoint /api/admin/scrape/dashboard returning the aggregates."),
    ("Activity stream tab",
     "Cross-job live event feed — every step row INSERT becomes an event. Color-coded by step_type, filterable by L1 / status / step. Polls /api/admin/scrape/activity?since=<timestamp>."),
    ("Companies-seed → scrape_jobs promotion endpoint",
     "API: POST /api/admin/scrape/candidates/[id]/promote. Creates a scrape_jobs row from the candidate and updates the candidate's promoted_to_job_id."),
    ("Per-L1 prompt tweaks",
     "Local Businesses needs service area + opening hours. Professional Services needs Bar / CPA certifications. Different prompt for each L1, organised at scripts/lib/prompts/<l1>.mjs."),
    ("Schema.org JSON-LD emission",
     "Phase 3 post-process: build a SoftwareApplication / Service / LocalBusiness JSON-LD object from the extracted JSON (Offer for pricing tiers, FAQPage for FAQs, AggregateRating if reviews). Inject into the /company/<slug> <head> for AEO."),
    ("Internal linking pass",
     "When an FAQ answer mentions a company name that has a /company/<slug> entry, replace the text with a link. Builds the internal graph for SEO."),
]
for title, detail in priorities:
    H3(title)
    P(detail)


# ═══════════════════════════════════════════════════════════════════
# END
# ═══════════════════════════════════════════════════════════════════
H1("End of context document")
P("Last updated automatically by scripts/write-scraper-doc.py.")
P(f"Output saved to: {OUT}")

doc.save(str(OUT))
print(f"Wrote {OUT}")
print(f"Size: {OUT.stat().st_size:,} bytes")

"""
Append Session 36 (May 9, 2026) to the running session report.

Mirrors S27-S35 structure: date/title line, Branch line, Executive
Summary, then numbered sections with prose paragraphs, a Files
Summary, Git Commits, and an end-of-report marker. No bullet
gimmicks — full sentences so future-Aadil can read it like a doc.
"""

from pathlib import Path
from docx import Document

DOC = Path(r"F:/infoWebWorld/InfoWebWorld_Session_Report_2026-03-23.docx")

doc = Document(str(DOC))


def H(text: str) -> None:
    doc.add_paragraph(text)


def P(text: str) -> None:
    doc.add_paragraph(text)


# ───────────────────────────────────────────────────────────────────────────
# Header
# ───────────────────────────────────────────────────────────────────────────
H("Session 36 — May 9, 2026")
H(
    "Session: Listing Verification Subsystem, Listing-Form Intro Flow + Step 2 "
    "Redesign + Edit-Page Chrome Compaction, Company-Listing Flow End-to-End "
    "(/profile/[slug] static page, smart hero gating, /dashboard/company, "
    "company → product autofill, sidebar cleanup, editorial company-page redesign)"
)
H("Developer: Claude Code (AI) + Aadil Parmar")
H("Branch: main (production)")

# ───────────────────────────────────────────────────────────────────────────
# Executive Summary
# ───────────────────────────────────────────────────────────────────────────
H("Executive Summary")
P(
    "By far the largest session yet — one continuous build covering five "
    "major shipped commits and ~5,000 net lines of code change. Began with "
    "the listing-verification subsystem (apply form + admin queue + public "
    "verified-by-InfoWebWorld badge + email notifications), pivoted to fix "
    "a regression in the admin reviews moderation route (missing "
    "updated_at column on the reviews table), then rebuilt the entry "
    "experience for /dashboard/new with two intro heroes (sector picker "
    "and company-vs-product picker) and a redesigned Step 2, then "
    "compacted the chrome around the entire form so the form card fills "
    "the available width, and finished with the marquee feature: a "
    "complete company-listing flow that lives in parallel to the existing "
    "product-listing flow with its own short 3-step form, its own static "
    "/profile/[slug] page, automatic cross-linking between products and "
    "companies, a clean slug rule (no more random suffixes), one-company-"
    "per-user enforcement, smart hero gating that retires the picker "
    "screens once they've been answered, one-time autofill of company "
    "fields onto a fresh product form, a focused /dashboard/company "
    "detail page, a full sidebar cleanup that stripped seven dead "
    "section headings, and a complete editorial redesign of the public "
    "company page after the first cut was rejected by Aadil for being "
    "too generic. Five commits, four migrations (one new), zero "
    "untouched corners of the dashboard."
)

# ───────────────────────────────────────────────────────────────────────────
# 1. Listing Verification Subsystem
# ───────────────────────────────────────────────────────────────────────────
H("1. Listing Verification Subsystem (commit 0698b47)")
P(
    "Built the owner-applied, admin-moderated identity-verification flow "
    "that earns a listing the prominent Verified-by-InfoWebWorld badge in "
    "its hero. Schema migration database/migration-listing-verification."
    "sql adds a listing_verification_requests audit-trail table (one row "
    "per application, many allowed per listing across re-applications, "
    "with legal_name / business_email / business_phone / "
    "registration_number / owner_role / document_url / social_proof JSON "
    "/ applicant_notes evidence fields, plus admin_notes / "
    "reviewed_by_admin_id / reviewed_at on the moderator side) and three "
    "resolved-state columns directly on submissions (verified TINYINT, "
    "verified_at DATETIME, verification_request_id BIGINT) so the read "
    "path on /company/[slug] never has to join verification_requests "
    "just to decide whether to render the badge."
)
P(
    "Owner-side at /dashboard/listings/[uuid]/verify: a state-aware "
    "single-page application form. When the listing is already verified "
    "the page renders a big mint card with the public-listing CTA and "
    "no form. When a request is pending review it shows an amber card "
    "with 'Submitted on {date}' plus a read-only summary of every field "
    "the applicant submitted. When the most recent request was rejected "
    "it shows the admin's note inline plus the form prefilled with the "
    "previous values for a quick fix-and-resubmit. When no request "
    "exists yet it's a fresh form. The form validates the business "
    "email's domain against the listing's website domain in real time "
    "and warns inline when they don't match (the strongest single "
    "ownership signal). On the /dashboard/listings page each card now "
    "shows a verification pill — Get Verified (coral) / Under review "
    "(amber) / Verified (mint with a shield) / Not verified · Re-apply "
    "(rejected with a dashed border) — plus an inline verified shield "
    "next to the company name when the listing is verified. The pill "
    "doubles as the entry-point — clicking it routes to the apply page."
)
P(
    "Admin-side at /iww-hq/verifications: a new dashboard nav entry "
    "leads to a two-pane queue mirroring /iww-hq/reviews. Each row "
    "shows the listing strip (logo + name + already-verified chip if "
    "applicable + submission date) on top, the applicant card "
    "(avatar + name + role) below, an evidence summary with a green "
    "Domain match or amber Domain mismatch flag on the business "
    "email row, and an expandable Show all evidence panel revealing "
    "any social profiles, document URL and applicant notes. Approve "
    "and Reject buttons live on the right with an optional admin-note "
    "textarea (recommended for rejects — included in the email to the "
    "owner). Approving flips submissions.verified to 1, sets "
    "verified_at to NOW(), points verification_request_id at the row, "
    "calls revalidatePath() against /company/[slug] so the badge goes "
    "live without waiting for the 48h auto-revalidate, and fires the "
    "owner email. Idempotent on re-approve clicks (no double email)."
)
P(
    "Email templates verificationApprovedEmail and "
    "verificationRejectedEmail were added to lib/email-templates.ts "
    "with matching dispatcher functions notifyOwnerOnVerificationApproved "
    "and notifyOwnerOnVerificationRejected in lib/notify-owner.ts. The "
    "approved email has a mint hero with a verified-shield avatar, the "
    "optional admin note inline, and a 'what this means for you' bullet "
    "list (prominent badge, higher trust signal, better conversion). "
    "The rejected email is amber, leads with the admin's reason, "
    "includes a 'common fixes' bullet list (matching email domain, "
    "registration number, owner role, public LinkedIn), and links "
    "directly to the verify page so they can re-apply."
)
P(
    "Public render: in app/listing/ListingDetailPage.tsx the existing "
    "always-on small shield next to the company name in the sticky "
    "head was gated on view.verified (preview mode keeps it on so the "
    "design preview at /test-listing-page shows the prominent variant). "
    "A new tlp-vbadge component sits between the tagline and the "
    "Last-updated line, rendering either a mint card with eyebrow / "
    "title / sub-copy + verified-on date when the listing is verified, "
    "or a quiet dashed-border pill with a shield-question icon reading "
    "'Unverified by InfoWebWorld' otherwise, so visitors can tell "
    "immediately whether a listing has completed verification. Both "
    "variants are 600px wide-max, sit just under the tagline, and "
    "respect the existing visual language of the listing page. The "
    "InfoBot floating chat widget was also hidden on /dashboard/* via a "
    "usePathname() check inside ChatWidget — the dashboard already has "
    "its own chrome and the launcher was competing visually."
)

# ───────────────────────────────────────────────────────────────────────────
# 2. Reviews Hotfix
# ───────────────────────────────────────────────────────────────────────────
H("2. Reviews Hotfix — Missing updated_at Column (commit c645f6f)")
P(
    "Aadil reported a server error when he visited /iww-hq/reviews to "
    "moderate a freshly-written review. Root cause: the admin "
    "moderation routes added on May 8 (commit 2542389) all reference "
    "reviews.updated_at — the GET /api/admin/reviews query selects it, "
    "and both POST /approve and POST /reject write `updated_at = NOW()`. "
    "But the original engagement-reviews migration (S35) created the "
    "reviews table without that column, so the moderation feature had "
    "been silently broken since the day it shipped (just nobody had "
    "moderated a review until now). New migration "
    "database/migration-reviews-updated-at.sql adds the column with "
    "ON UPDATE CURRENT_TIMESTAMP. Original migration-engagement-reviews."
    ".sql was also updated for fresh installs with a header note "
    "pointing existing installs at the hotfix file."
)

# ───────────────────────────────────────────────────────────────────────────
# 3. Listing-Form Intro Flow + Step 2 Redesign + Chrome Compaction
# ───────────────────────────────────────────────────────────────────────────
H("3. Listing-Form Intro Flow + Step 2 Redesign + Chrome Compaction (commit fbdd287)")
P(
    "Major UX overhaul of /dashboard/new and the listing form. Split "
    "into three threads."
)
P(
    "Thread A — Two intro heroes before the rail/form. First "
    "SectorPickHero asks 'Where does your business live?' with six "
    "simple text pills (Artificial Intelligence & ML / IT Services & "
    "Agencies / Local Businesses / Professional Services / Software & "
    "SaaS / Startups & Innovation), centered title and subhead, big "
    "Clash-Royale-style mascot anchored bottom-right at clamp("
    "460-720px) wide and clamp(520-780px) tall. Pure pill shape — "
    "round border, sector name only, no dots / icons / accent washes "
    "— matching the rest of the dashboard's design language. Hover "
    "fills the pill black; pick fills it coral and fades the siblings "
    "to 35% as the hero unmounts. The mascot loads from "
    "public/illustrations/welcome-mascot.png with a fail-quiet "
    "onError so a missing asset never produces a broken-image icon. "
    "Then the second hero ListingModePickHero asks 'What are you "
    "listing?' with two pills whose labels and subhead text adapt "
    "per-sector: AI/ML pills read 'AI Company' / 'AI Tool', Software "
    "& SaaS reads 'Software Company' / 'Product', IT Services reads "
    "'Agency' / 'Service', Startups reads 'Startup' / 'Product', "
    "Local Businesses reads 'Business' / 'Service', Professional "
    "Services reads 'Firm' / 'Service'. The mascot for this hero is "
    "a different pose (thumbs-up / pointing) and lives bottom-LEFT so "
    "the two heroes don't feel identical. Underlying value stays "
    "'company' or 'product' regardless of label — only the displayed "
    "wording changes — so downstream form code doesn't have to "
    "branch on every sector. Both heroes hide entirely on edit mode. "
    "Draft restoration strips l1Id / l2Id / l3Id / listingMode from "
    "localStorage so each fresh /dashboard/new visit re-runs them; "
    "the rest of the draft (tagline, description, etc.) still comes "
    "back."
)
P(
    "Thread B — Step 2 (Category & Classification) redesigned from "
    "three flat Field rows into two card blocks. Block 1 'Where does "
    "your listing live?' contains the 3-level category picker plus a "
    "live breadcrumb summary that appears once L1 is picked ('Software "
    "& SaaS › AI / ML › Computer Vision'), a 3-dot progress indicator "
    "in the header, and Specializations as a sub-section inside the "
    "same card (they're contextual to L3, not a sibling concept) with "
    "a counter showing '2 of 8 selected'. Block 2 'How would you "
    "describe your business?' renders the 5 tag groups as collapsible "
    "accordion rows. Each row's header has a colored dot in the "
    "group's color, the group name, a count chip ('Select' or '3 "
    "selected') filled in the group's accent, and a chevron. The "
    "Location tag group was hidden from the form entirely since "
    "geographic context is already captured precisely on Step 3 — was "
    "a duplicate ask. Filter applied at the load site so render and "
    "validation both ignore it. The card header shows a 'X of N done' "
    "coverage chip that turns mint when complete, giving the user "
    "obvious progress feedback. First group is expanded by default."
)
P(
    "Thread C — Edit page (and /dashboard/new) chrome was eating "
    "real estate. Aadil circled the empty space on the left and right "
    "of the form card, plus the entire header bar at the top, asking "
    "for it all gone. Fix scoped via :has(.df-wrap) so only form "
    "pages compact while the rest of the dashboard keeps its "
    "generous spacing. Removed the 1100px max-width clamp on .nl--form "
    "(was wasting ~250px on each side on big screens). Hid the "
    ".tp-page-head (title / breadcrumb / icons / avatar — duplicates "
    "of what the dashboard sidebar already shows) and the "
    ".df-plan-strip (FREE / $0 / Draft saved bar) entirely on form "
    "pages. Page padding dropped from 28px / clamp(24-44px) / 40px to "
    "6px / 6px / 6px. Form wrap gap 14 → 10. Grid gap 18 → 12. Rail "
    "width 220 → 188. Body padding 20px 24px → 14px 14px. Rail and "
    "form-card borders bumped from 1px var(--df-border) to 1.5px "
    "var(--df-border-dark) so the structure reads as deliberate. The "
    "rail no longer has align-self: start or position: sticky (the "
    "page itself doesn't scroll on form pages — the form card "
    "manages its own internal scroll), so it stretches to match the "
    "form card's height. The 'Your plan' card we'd added to the rail "
    "moves to the bottom via margin-top: auto so empty space falls "
    "between the step list and the plan card / progress strip — looks "
    "intentional, not floating. All field labels were rewritten from "
    "11px UPPERCASE letter-spacing-tracked to 14px sentence case 700 "
    "weight (.df-label, .df-cat-title, .df-step2-spec-label) — bigger, "
    "bolder, more readable; hierarchy preserved (section title 18px > "
    "card title 15px > field label 14px > sub-labels 13px)."
)

# ───────────────────────────────────────────────────────────────────────────
# 4. Company-Listing Flow End-to-End
# ───────────────────────────────────────────────────────────────────────────
H("4. Company-Listing Flow End-to-End (commit b4f554c)")
P(
    "The marquee feature: a parallel listing flow for COMPANIES "
    "alongside the existing PRODUCT flow. A user can have one "
    "company profile (one-per-user) plus many product listings; "
    "products auto-link to the user's company so every product page "
    "surfaces 'Made by {Company}' pointing to the company profile, "
    "and every company profile lists 'Products by us' pointing back."
)
P(
    "Schema (database/migration-listings-company-mode.sql): one ALTER "
    "TABLE on submissions adds listing_mode ENUM('product','company') "
    "DEFAULT 'product' (so existing rows are unchanged), "
    "parent_company_id BIGINT UNSIGNED NULL self-referential FK on "
    "delete set null, is_hiring TINYINT, plus indexes on listing_mode "
    "and parent_company_id. One unified table keeps the admin "
    "moderation queue, approval flow, and deploy hook all working "
    "without a parallel companies table to maintain."
)
P(
    "Slug rule change for NEW listings only — existing slugs untouched. "
    "POST /api/submissions no longer appends the random "
    "`-${uuid.slice(0,8)}` suffix. Slug is now pure slugify(name). "
    "Uniqueness is enforced per listing_mode (companies live at "
    "/profile/<slug> and products at /company/<slug> — different URL "
    "namespaces, can share slug strings). On collision the API rejects "
    "with code SLUG_TAKEN and a clear message instructing the user to "
    "pick a different name (Twitter / GitHub-handle style). One-"
    "company-per-user is also enforced at the API layer: a second POST "
    "with listingMode='company' from the same authed user rejects with "
    "code COMPANY_EXISTS and the slug of the existing one so the "
    "client can deep-link them to edit."
)
P(
    "Company form is a focused 3-step short flow distinct from the "
    "long 7-step product form: Step 1 Identity (logo, name, tagline, "
    "website, description, founded year, team size, header tags), "
    "Step 2 Details (contact name, email, phone, country/state/city, "
    "HQ address, LinkedIn / Twitter / Facebook, Hiring toggle), Step "
    "3 Review (read-only summary + Submit profile button). New step "
    "components live in app/dashboard/new/form/steps/CompanyStep[1-3]"
    "*.tsx and reuse all existing form components (Field, Uploader, "
    "ChipInput, PhoneRow, Select). DashboardListingForm picks "
    "COMPANY_STEPS vs PRODUCT_STEPS based on form.listingMode, set by "
    "the listing-mode hero earlier in the flow. Submit button reads "
    "'Submit profile' in company mode."
)
P(
    "Public company page at /profile/[slug] uses the same static SSG "
    "delivery model as /company/[slug]: revalidate=172800, "
    "dynamicParams=false, generateStaticParams pulls every active/"
    "paid company at build time. JSON-LD Organization schema with "
    "foundingDate, employees, address, email, telephone, sameAs "
    "(website + linkedin + twitter + facebook). The first cut had "
    "eight sections (sticky head / title block / About / Products by "
    "us / Where we are / Team & culture / Get in touch / FAQs) and "
    "Aadil shot it down for being too generic and corporate — the "
    "redesign comes later in the session."
)
P(
    "Cross-links: in app/listing/ListingDetailPage.tsx the product "
    "page hero now shows a quiet 'Made by {Company} ↗' pill linking "
    "to /profile/[slug] when parent_company_id is set on the row. "
    "Fetched server-side via a small follow-up queryOne in "
    "/company/[slug]/page.tsx; tolerated when the migration hasn't "
    "run. POST /api/submissions auto-links every new product listing "
    "to the authed user's approved company profile (no UI needed in "
    "the product form — there's only ever one company per user). The "
    "/api/admin/listings/[slug]/revalidate route was updated to read "
    "the row's listing_mode and revalidate /profile/[slug] for "
    "company rows or /company/[slug] for product rows so the deploy "
    "hook hits the right path."
)
P(
    "Admin /iww-hq/submissions got a small filter chip row beneath the "
    "status tabs: All / Products / Companies, with live counts. "
    "Existing approve / reject / deploy actions work for both modes "
    "unchanged (one table, one moderation flow). RealSubmission "
    "gained listingMode and parentCompanyId fields, read by mapRow "
    "with safe defaults for legacy rows. /dashboard/listings now "
    "splits into two visual sections: 'Your company profile' (single "
    "coral-bordered card with logo, name, tagline, status, Edit / "
    "View profile actions) at the top, and 'Your products' (existing "
    "grid filtered to listing_mode != 'company') below. The section "
    "heading for products only appears when a company also exists, "
    "so a user with only products sees the unmodified one-section "
    "view."
)
P(
    "Defensive pre-migration tolerance: every new query that "
    "references listing_mode or parent_company_id is wrapped in a "
    "try/catch that detects the 'Unknown column' error and falls back "
    "to legacy SELECTs that omit the new columns. So the app keeps "
    "rendering before phpMyAdmin runs the migration, treating every "
    "row as listing_mode='product' for the duration. The first "
    "version of the dashboard listings fallback had a bug: when "
    "MySQL only reported the FIRST missing column (verified) the "
    "fallback still tried to use COALESCE(s.listing_mode) and broke. "
    "Fix was to always drop down to LEGACY_BASE on either-error so we "
    "survive any combination of pending migrations."
)

# ───────────────────────────────────────────────────────────────────────────
# 5. Smart Hero Gating + Sidebar Cleanup + /dashboard/company + Editorial Profile Redesign
# ───────────────────────────────────────────────────────────────────────────
H("5. Smart Hero Gating + Sidebar Cleanup + /dashboard/company + Editorial Profile Redesign (commit 889e420)")
P(
    "Five-part overhaul tying the company-listing flow into the rest "
    "of the dashboard."
)
P(
    "Sidebar cleanup: stripped seven dead-section nav headings "
    "(Business Listing / Discovery & Visibility / Lead Management / "
    "Reviews & Reputation / Community / Analytics & Insights / "
    "Support & Admin) — every page they pointed to was placeholder "
    "content. The featureNav useMemo array was removed from "
    "DashboardShell.tsx's render block (the SECTIONS / FEATURES / "
    "findFeature constants are kept since they're still used by /plans "
    "and the section/feature routes, just no longer rendered in the "
    "sidebar nav). New 'My Company' entry conditionally inserted "
    "between My Listings and New Listing when the user has a "
    "company row in submissions (any status). Sourced from a "
    "hasCompany boolean prop set server-side in app/dashboard/layout."
    "tsx via a small queryOne wrapped in try/catch (pre-migration "
    "tolerance again)."
)
P(
    "Smart hero gating: server pre-fetches the user's most-recent "
    "listing in app/dashboard/new/page.tsx and passes an "
    "existingContext snapshot ({ hasAny, hasCompany, l1Id, "
    "companyPrefill }) through NewListingClient into "
    "DashboardListingForm. The hero gates derive from it — any prior "
    "listing exists means sector hero is skipped (l1Id locked from "
    "the prior listing's L1, sourced by walking the category parent "
    "chain up to 5 hops); company listing exists means mode hero is "
    "skipped too (listingMode auto-locks to 'product' since company "
    "is one-per-user); only product listings exist means sector "
    "skipped but mode hero still shown so the user can finally do "
    "their company; nothing exists means both heroes show as before. "
    "The rule matches Aadil's reasoning: an AI company isn't going to "
    "switch to local biz, but a user who's only made products might "
    "still want to claim their company profile."
)
P(
    "One-time autofill from company → product form: when a user with "
    "an approved company starts a new product listing, the lazy state "
    "initializer in DashboardListingForm seeds these reusable fields "
    "from the company row: contactName, email, phoneCode, phone, "
    "countryCode, country, state, city, hqLocation, linkedin, "
    "twitter, facebook, founded, employees. Product-specific fields "
    "(companyName, tagline, description, logo, headerTags, website) "
    "stay empty. Pure one-time seed; the user can edit anything and "
    "the localStorage draft restoration runs after, so once edited "
    "the change wins forever — the strip-list (l1Id/l2Id/l3Id/"
    "listingMode) doesn't include these fields. The prefill data is "
    "fetched server-side in /dashboard/new/page.tsx via a queryOne "
    "with a country join, so countryCode comes through as the ISO "
    "code (countries.code) rather than the numeric country_id."
)
P(
    "New /dashboard/company page (app/dashboard/company/page.tsx) "
    "as a focused detail view of the user's single company profile. "
    "Identity strip with logo + name + tagline + status pill + "
    "Verified-by-IWW chip when applicable + Hiring chip when "
    "applicable + Edit profile / View public profile buttons. 4-tile "
    "stats grid showing Products linked (with live/pending split), "
    "Team size + founded year, Plan + creation date, Verified-by-"
    "InfoWebWorld status with a Get-verified or Re-apply CTA when "
    "appropriate. 'Company information' read-only block with EVERY "
    "captured field rendered as a key-value row (description, "
    "founded, team size, funding, hiring, headquarters, country, "
    "state, city, contact name, contact email, phone, LinkedIn, "
    "Twitter, Facebook). 'Products under {Company}' grid linking to "
    "each product's edit page + view public listing. Three-level "
    "pre-migration fallback in the queryOne — if listing_mode column "
    "is missing the page redirects to /dashboard/listings; if "
    "verified column is missing it drops those fields and renders "
    "anyway; full query otherwise. CSS lives in app/styles/dashboard."
    "css under the .dco-* namespace."
)
P(
    "/profile/[slug] complete editorial redesign — the senior "
    "designer feedback on the first cut was that it was too generic "
    "and corporate, like a 2018 startup landing page. Rebuilt as four "
    "focused sections distinct from the dense product detail page. "
    "Hero is a soft sector-tinted backdrop (radial gradient using the "
    "L1 sector's color from categories.color, set as --cmp-accent on "
    "the React root so any descendant can pick it up via var() "
    "without prop-drilling), with a large 132px logo, a big "
    "clamp(34-56px) company name with a verified shield inline next "
    "to it, an eyebrow showing the sector, the tagline, header-tag "
    "chips, and a single-line dot-separated meta strip showing "
    "Founded · Team size · Location · Hiring (replaces the previous "
    "stat-tile-soup of 4 separate cards). Two CTAs — Visit website "
    "filled in the sector accent + outline Get-a-quote. About is a "
    "65/35 split: longform description paragraphs (16.5px, line-"
    "height 1.75) on the left rendered as readable typography with "
    "blank-line paragraph breaks; supporting card stack on the right "
    "with website link (sector-accent border), social icons row, and "
    "verification status as supporting detail (not crowding the "
    "hero). Portfolio is the centerpiece: 'Products by {Company}' "
    "with a count chip filled in the sector accent. 2-column grid "
    "of large rich product cards (no chevron clutter): 56px logo top-"
    "left + category pill top-right, big 22px name, 2-line clamped "
    "tagline, divider line, starting price + accent-colored View "
    "Product link. Each card animates a 3px sector-accent stripe at "
    "the top on hover plus a 12px sector-accent border + 3px lift. "
    "Empty state is a wide dashed card inviting the company to add "
    "products. Connect is a single block with no card chrome — copy "
    "+ actions row (Get a quote primary, website outline) + a "
    "divider line + meta row with email/phone inline links and "
    "social icons pushed to the right. Sections explicitly removed "
    "vs the previous version: separate Where-we-are (folded into "
    "hero meta), Team-and-culture placeholder (no real data), and "
    "the FAQ section (no source). Less is more — four sections vs "
    "the previous eight. The L1 sector color drives the accent so an "
    "AI/ML company reads as electric blue, an IT agency as teal, a "
    "startup as violet — quietly differentiated without per-company "
    "branding work. CSS lives in app/styles/profile-page.css with "
    "the .cmp-* namespace, full rewrite from the first version."
)

# ───────────────────────────────────────────────────────────────────────────
# 6. Files Modified
# ───────────────────────────────────────────────────────────────────────────
H("6. Files Modified (Across All Five Commits)")
P(
    "Database migrations (run in phpMyAdmin against "
    "cdbrisgy_infowebworld): "
    "database/migration-listing-verification.sql (NEW — "
    "listing_verification_requests table + verified columns on "
    "submissions). "
    "database/migration-reviews-updated-at.sql (NEW — hotfix for "
    "missing reviews.updated_at column). "
    "database/migration-listings-company-mode.sql (NEW — "
    "listing_mode + parent_company_id + is_hiring on submissions). "
    "Also updated database/migration-engagement-reviews.sql for "
    "fresh installs to include the updated_at column from the start."
)
P(
    "API routes (new): app/api/dashboard/listings/[uuid]/verify/"
    "route.ts (owner submits verification application); app/api/"
    "admin/verifications/route.ts (admin GET list by status); app/"
    "api/admin/verifications/[id]/approve/route.ts and reject/route."
    "ts (admin moderation actions, approve flips submissions."
    "verified=1 and revalidates /company/[slug] and emails owner)."
)
P(
    "API routes (modified): app/api/submissions/route.ts (slug rule "
    "rewrite, listingMode handling, one-company-per-user enforcement, "
    "parent_company_id auto-link, full INSERT extended with new "
    "columns); app/api/admin/listings/[slug]/revalidate/route.ts "
    "(mode-aware path resolution: /profile/[slug] for company rows "
    "or /company/[slug] for product rows, with pre-migration "
    "fallback)."
)
P(
    "Dashboard pages and components (new): "
    "app/dashboard/listings/[uuid]/verify/page.tsx + VerifyApplyClient."
    "tsx (verification apply flow with state-aware UI); "
    "app/dashboard/new/form/components/SectorPickHero.tsx (entry "
    "splash with 6 sector pills + bottom-right mascot); "
    "app/dashboard/new/form/components/ListingModePickHero.tsx "
    "(second splash with 2 mode pills + bottom-left mascot, sector-"
    "adaptive labels); "
    "app/dashboard/new/form/steps/CompanyStep1Identity.tsx, "
    "CompanyStep2Details.tsx, CompanyStep3Review.tsx (3-step "
    "company form); "
    "app/dashboard/company/page.tsx (My Company detail page with "
    "stats, info readout, products grid)."
)
P(
    "Dashboard pages and components (modified): "
    "app/dashboard/DashboardShell.tsx (sidebar cleanup — featureNav "
    "removed, conditional My Company entry); "
    "app/dashboard/layout.tsx (server-fetches hasCompany flag); "
    "app/dashboard/listings/page.tsx (verification pill + Get "
    "Verified CTA + company section split with three-level pre-"
    "migration fallback); "
    "app/dashboard/new/page.tsx (server pre-fetches existingContext "
    "with category-chain walk); "
    "app/dashboard/new/NewListingClient.tsx (forwards prop); "
    "app/dashboard/new/form/DashboardListingForm.tsx (mode-aware "
    "step rendering, hero gates, autofill seeding, mode-aware submit "
    "payload with field zeroing for company mode); "
    "app/dashboard/new/form/components/RailNav.tsx (in-rail Your "
    "Plan card with per-tier accent); "
    "app/dashboard/new/form/steps/Step2Category.tsx (full rewrite "
    "as two card blocks with collapsible tag-group accordions); "
    "app/dashboard/new/form/types.ts (listingMode + isHiring + "
    "ListingMode union); "
    "app/dashboard/new/form/constants.ts (PRODUCT_STEPS / "
    "COMPANY_STEPS arrays, isHiring default); "
    "app/dashboard/new/form/validation.ts (company-mode step cases)."
)
P(
    "Public listing pages (modified): "
    "app/company/[slug]/page.tsx (parent-company linkage join with "
    "pre-migration fallback); "
    "app/listing/ListingDetailPage.tsx (gated existing verified "
    "shield, new tlp-vbadge mint/dashed pill, Made-by-Company pill "
    "linking to /profile)."
)
P(
    "Public profile page (new): "
    "app/profile/[slug]/page.tsx (server SSG with category color "
    "join + JSON-LD Organization); "
    "app/profile/CompanyDetailPage.tsx (client component, four-"
    "section editorial design, full rewrite from the first cut)."
)
P(
    "Admin pages (modified): "
    "app/iww-hq/components/AdminShell.tsx (Verifications nav entry); "
    "app/iww-hq/verifications/page.tsx (NEW — two-pane moderation "
    "queue with evidence summaries, domain-match flag, approve/reject "
    "with admin note); "
    "app/iww-hq/submissions/page.tsx (Companies/Products filter chip "
    "row + listingMode+parentCompanyId on RealSubmission); "
    "app/iww-hq/submissions/submissions.css (.sub-mode-chip styles); "
    "app/iww-hq/data/submissions-storage.ts (verified, verifiedAt, "
    "listingMode, parentCompanyId on RealSubmission with safe "
    "defaults via mapRow)."
)
P(
    "Email + notify (modified): "
    "lib/email-templates.ts (verificationApprovedEmail, "
    "verificationRejectedEmail); "
    "lib/notify-owner.ts (notifyOwnerOnVerificationApproved, "
    "notifyOwnerOnVerificationRejected dispatchers)."
)
P(
    "Styles (modified): "
    "app/styles/dashboard.css (verification pills, dash-co-* company "
    "section card on /dashboard/listings, dco-* /dashboard/company "
    "page styles); "
    "app/styles/dashboard-form.css (hero spk-* + step2 spk- accordion "
    "+ chrome compaction overrides + sentence-case labels + plan-card "
    "in rail + company-form .df-toggle-row + .df-rev-* compact review "
    "styles); "
    "app/styles/dashboard-shell.css (:has(.df-wrap) chrome compaction "
    "scoped overrides — page padding, header hide, plan strip hide, "
    "icon button shrinks); "
    "app/styles/test-listing-page.css (.tlp-vbadge mint/dashed "
    "verified card + .tlp-made-by linkage pill); "
    "app/styles/profile-page.css (NEW + REWRITTEN — full editorial "
    ".cmp-* design system); "
    "app/components/chat/ChatWidget.tsx (usePathname dashboard hide); "
    "app/globals.css (profile-page.css import added)."
)
P(
    "Public assets (new): "
    "public/illustrations/welcome-mascot.png (sector hero — Clash "
    "builder waving + pointing right, 7.2 MB); "
    "public/illustrations/welcome-mascot-2.png (mode hero — same "
    "character thumbs-up + pointing left, 10.7 MB)."
)

# ───────────────────────────────────────────────────────────────────────────
# 7. Git Commits (Chronological)
# ───────────────────────────────────────────────────────────────────────────
H("7. Git Commits (Chronological)")
P(
    "0698b47 — feat: listing verification subsystem + hide InfoBot on "
    "dashboard. 18 files / +2510 / -12. Schema migration, apply form, "
    "admin queue, email templates, public badge, ChatWidget pathname "
    "guard."
)
P(
    "c645f6f — fix(reviews): add missing updated_at column to reviews "
    "table. 2 files / +25 / -0. Hotfix migration + canonical "
    "migration update for fresh installs."
)
P(
    "fbdd287 — feat(dashboard): listing-form intro heroes + Step 2 "
    "redesign + compact chrome. 11 files / +1155 / -89. Both intro "
    "heroes with adaptive labels, Step 2 two-card redesign with "
    "collapsible tag groups, sector accent dropped per Aadil's "
    "feedback, page+form chrome compaction scoped via :has(.df-wrap), "
    "1.5px dark borders, plan card in rail, sentence-case 14px field "
    "labels."
)
P(
    "b4f554c — feat: company listings — short form, /profile/[slug] "
    "page, cross-links. 25 files / +2142 / -84. Schema migration, "
    "3-step company form, public /profile/[slug] static SSG (first "
    "cut, 8 sections — later replaced), Made-by-Company cross-link "
    "on product pages, auto-link on submit, admin Companies/Products "
    "filter, dashboard split."
)
P(
    "889e420 — feat(dashboard): /dashboard/company + smart hero "
    "gating + sidebar cleanup, redesigned /profile. 9 files / +1566 / "
    "-530. Sidebar dead-headings stripped, conditional My Company "
    "entry, smart hero gating from existingContext, one-time autofill "
    "company → product, /dashboard/company detail page, full "
    "editorial redesign of /profile/[slug] with sector-accent system."
)

# ───────────────────────────────────────────────────────────────────────────
# 8. Decisions Captured
# ───────────────────────────────────────────────────────────────────────────
H("8. Decisions Captured")
P(
    "URL strategy. /company/[slug] stays as the product page (zero "
    "migration risk for existing live URLs). /profile/[slug] is the "
    "new namespace for company pages. Plural / singular distinguishes "
    "without breaking shares or SEO indexing on existing product "
    "URLs. Future could rename product pages to /product/[slug] but "
    "not now."
)
P(
    "Slug rule. Going forward, slug = slugify(name) with NO suffix. "
    "Collision = hard reject (Twitter / GitHub style). Existing rows "
    "keep their gemini-ai-715bb01d-style suffixes — only new "
    "submissions get clean URLs. Per-listing-mode uniqueness so a "
    "company at /profile/openai and a product at /company/openai can "
    "coexist if their slugs happen to match (different namespaces)."
)
P(
    "One-company-per-user. Enforced at the API layer rather than via "
    "a partial UNIQUE index since partial indexes aren't portable "
    "across MySQL versions. The API check is the right place for "
    "that rule anyway."
)
P(
    "One submissions table for both modes. listing_mode column "
    "distinguishes. Same admin queue, approval flow, deploy hook, "
    "moderation tooling — no parallel companies table to maintain. "
    "Indexes on listing_mode and parent_company_id keep query plans "
    "tight."
)
P(
    "Pre-migration tolerance everywhere. Every new query that "
    "references listing_mode / parent_company_id / verified / "
    "verified_at / listing_verification_requests is wrapped in a "
    "try/catch detecting Unknown column and falling back to a "
    "legacy SELECT. The app keeps rendering before phpMyAdmin runs "
    "the migration; only the new feature is invisible until then. "
    "Fallback bug: MySQL only reports the FIRST missing column, so "
    "fallbacks must drop ALL new columns when ANY is detected "
    "missing — not just the one named in the error."
)
P(
    "Auto-link products to company at submit time, not via UI. The "
    "product form has no 'Which company makes this?' picker since "
    "there's only ever one company per user. POST /api/submissions "
    "auto-attaches authedUser's company to every new product. Clean "
    "implementation, zero extra clicks for the user."
)
P(
    "One-time autofill, not always-on linking. When a product form "
    "mounts and the user has an approved company, prefill reusable "
    "fields (contact, location, socials, founded, team size). Once "
    "the user edits a field, the change wins forever (draft "
    "autosave doesn't strip those fields). Product-specific fields "
    "(companyName, tagline, description, logo, headerTags, website) "
    "are NEVER pre-filled — they're per-product copy and would be "
    "wrong shared across products."
)
P(
    "Smart hero gating, not unconditional re-shows. After ANY "
    "listing exists the sector pick is locked permanently from the "
    "user's first listing's L1 — an AI company isn't switching to "
    "local biz. After a COMPANY listing exists the mode pick is "
    "also locked to product (one-company-per-user). After only "
    "PRODUCT listings the mode hero still shows so the user can "
    "still claim their company. Heroes only show in full when the "
    "user has zero listings. Edge case: rejected-only listings still "
    "count as 'something exists in the account' for hero-skipping "
    "purposes."
)
P(
    "Editorial company-page design over a generic stat-tile layout. "
    "First cut was rejected by the senior designer eye for being too "
    "corporate. Rebuilt as four focused sections with the sector "
    "color as a quiet through-line and Products as the visual "
    "centerpiece — companies exist on this site to showcase their "
    "products. Stats live as a single-line meta strip in the hero, "
    "not stat tiles. Where-we-are / Team / FAQs sections were "
    "deleted entirely; the data they showed either lives in the "
    "hero meta strip now (location, hiring) or didn't have a real "
    "source (team page, FAQs)."
)
P(
    "Sidebar shows only what works. Removed seven dead-section "
    "headings (Business Listing through Support & Admin) — every "
    "page they pointed at was placeholder content and surfacing "
    "them looked like the dashboard had way more functionality "
    "than it does. Sidebar now: Home / My Listings / [My Company "
    "if has one] / New Listing / Settings / Browse Site. SECTIONS "
    "constants kept because /plans + admin still use them, just no "
    "longer rendered in the sidebar."
)
P(
    "InfoBot off on /dashboard/*. The floating chat launcher is "
    "useful for visitors browsing the public site but conflicts "
    "with the dashboard's own chrome. Hidden via a usePathname() "
    "early-return inside ChatWidget."
)

# ───────────────────────────────────────────────────────────────────────────
# 9. Pending Migrations as of Session End
# ───────────────────────────────────────────────────────────────────────────
H("9. Migrations Status as of Session End (May 9, 2026 EOD)")
P(
    "All four pre-existing pending migrations from S34/S35/May 8 had "
    "already been run before this session began (per Aadil's "
    "confirmation): migration-listings-v3.sql, migration-user-plan-"
    "purchases.sql, migration-engagement-reviews.sql, migration-leads-"
    "fields.sql. New migrations created during S36: migration-"
    "listing-verification.sql (verification subsystem — Aadil "
    "confirmed run), migration-reviews-updated-at.sql (reviews hotfix "
    "— Aadil confirmed run), migration-listings-company-mode.sql "
    "(company-listing flow — Aadil confirmed run, fixed the runtime "
    "error he hit on /dashboard/listings)."
)
P(
    "Notable late-session bug: after Aadil confirmed all migrations "
    "run, /dashboard/company hit a server error 'Unknown column "
    "s.verified'. The verification migration evidently hadn't fully "
    "applied or had been partially rolled back. Defensive 3-level "
    "fallback was added to /dashboard/company — full query first, "
    "then drop verified columns on Unknown-column-verified, then "
    "redirect on Unknown-column-listing_mode — so the page renders "
    "regardless of which migration combination is actually live."
)

H("— End of Session 36 Report —")

doc.save(str(DOC))
print(f"Saved Session 36 to {DOC}")
